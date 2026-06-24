# """
# backend.py
# ===========
# Core RAG (Retrieval-Augmented Generation) backend for the Qwen3B chatbot.

# Responsibilities:
# - Load and extract text from uploaded documents (PDF, scanned PDF/image via OCR,
#   DOCX, CSV, XLSX, TXT)
# - Chunk and embed document text, store/retrieve via a local FAISS vector index
# - Perform optional web search (DuckDuckGo) and scrape (BeautifulSoup) for live
#   information
# - Build a final prompt combining chat history + retrieved context, and call the
#   Qwen3B model hosted on Kaggle (exposed via an ngrok HTTP endpoint) to generate
#   a response

# This module has NO Streamlit dependency - it is pure backend logic so it can be
# tested or reused independently of the UI.
# """

# import os
# import logging
# from typing import Any, Dict, List, Mapping, Optional, Tuple

# import requests
# import pandas as pd
# from pypdf import PdfReader
# import pytesseract
# from pdf2image import convert_from_path
# from PIL import Image
# import docx  # python-docx
# from bs4 import BeautifulSoup

# from langchain_core.language_models.llms import LLM
# from langchain_core.callbacks.manager import CallbackManagerForLLMRun
# from langchain_core.documents import Document
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# from langchain_huggingface import HuggingFaceEmbeddings

# # The duckduckgo-search package was renamed to "ddgs". Support both so this
# # keeps working regardless of which one is installed.
# try:
#     from ddgs import DDGS
# except ImportError:
#     from duckduckgo_search import DDGS


# # ---------------------------------------------------------------------------
# # Logging configuration
# # ---------------------------------------------------------------------------
# logging.basicConfig(
#     level=logging.INFO,
#     format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
#     handlers=[
#         logging.StreamHandler(),
#         logging.FileHandler("rag_chatbot.log", encoding="utf-8"),
#     ],
# )
# logger = logging.getLogger("rag_backend")


# # ---------------------------------------------------------------------------
# # Configuration
# # ---------------------------------------------------------------------------
# class Config:
#     """
#     Central configuration for the backend.
#     NGROK_URL and APP_API_KEY can be overridden via environment variables so
#     you don't have to edit code every time the Kaggle notebook restarts and
#     issues a new ngrok URL.
#     """

#     NGROK_URL: str = os.getenv("NGROK_URL", "https://unshakable-yasuko-luxuriantly.ngrok-free.dev")
#     APP_API_KEY: str = os.getenv("APP_API_KEY", "qwen-server-prod-2026-7f9a4c2d8e1b5f6a")

#     EMBEDDING_MODEL_NAME: str = "sentence-transformers/all-MiniLM-L6-v2"
#     VECTORSTORE_DIR: str = "faiss_index"

#     CHUNK_SIZE: int = 800
#     CHUNK_OVERLAP: int = 100
#     TOP_K: int = 4

#     # If a PDF page has fewer than this many extracted characters on average,
#     # we assume it's a scanned/image-based PDF and fall back to OCR.
#     OCR_TEXT_LENGTH_THRESHOLD: int = 20

#     MAX_NEW_TOKENS: int = 512
#     TEMPERATURE: float = 0.2

#     WEB_SEARCH_MAX_RESULTS: int = 3
#     SCRAPE_TIMEOUT: int = 10
#     SCRAPE_MAX_CHARS: int = 3000
#     LLM_REQUEST_TIMEOUT: int = 120


# # ---------------------------------------------------------------------------
# # Custom LangChain LLM wrapper around the Kaggle / ngrok endpoint
# # ---------------------------------------------------------------------------
# class QwenNgrokLLM(LLM):
#     """
#     LangChain-compatible LLM that calls the FastAPI /generate endpoint running
#     on Kaggle and exposed publicly via ngrok. The endpoint is NOT OpenAI-API
#     compatible, so this wrapper translates LangChain's standard call interface
#     into the custom {"prompt", "max_new_tokens", "temperature"} request schema.
#     """

#     ngrok_url: str
#     api_key: str
#     max_new_tokens: int = Config.MAX_NEW_TOKENS
#     temperature: float = Config.TEMPERATURE
#     timeout: int = Config.LLM_REQUEST_TIMEOUT

#     @property
#     def _llm_type(self) -> str:
#         return "qwen-ngrok"

#     def _call(
#         self,
#         prompt: str,
#         stop: Optional[List[str]] = None,
#         run_manager: Optional[CallbackManagerForLLMRun] = None,
#         **kwargs: Any,
#     ) -> str:
#         logger.info("Calling Qwen LLM endpoint (prompt length=%d chars)", len(prompt))
#         try:
#             response = requests.post(
#                 f"{self.ngrok_url}/generate",
#                 json={
#                     "prompt": prompt,
#                     "max_new_tokens": kwargs.get("max_new_tokens", self.max_new_tokens),
#                     "temperature": kwargs.get("temperature", self.temperature),
#                 },
#                 headers={"x-api-key": self.api_key},
#                 timeout=self.timeout,
#             )
#             response.raise_for_status()
#             text = response.json()["response"]
#             logger.info("LLM responded with %d characters", len(text))
#             return text
#         except requests.exceptions.RequestException as e:
#             logger.error("LLM request failed: %s", e)
#             return f"[Error contacting the model: {e}]"

#     @property
#     def _identifying_params(self) -> Mapping[str, Any]:
#         return {"ngrok_url": self.ngrok_url, "max_new_tokens": self.max_new_tokens}


# def create_llm() -> QwenNgrokLLM:
#     """Factory to create the Qwen LLM client using values from Config."""
#     return QwenNgrokLLM(
#         ngrok_url=Config.NGROK_URL,
#         api_key=Config.APP_API_KEY,
#         max_new_tokens=Config.MAX_NEW_TOKENS,
#         temperature=Config.TEMPERATURE,
#     )


# def check_llm_health() -> bool:
#     """Ping the /health endpoint of the ngrok-exposed Kaggle server."""
#     try:
#         response = requests.get(f"{Config.NGROK_URL}/health", timeout=10)
#         return response.status_code == 200
#     except requests.exceptions.RequestException as e:
#         logger.warning("LLM health check failed: %s", e)
#         return False


# # ---------------------------------------------------------------------------
# # Document loading & text extraction
# # ---------------------------------------------------------------------------
# def extract_text_from_pdf(file_path: str) -> str:
#     """
#     Extract text from a PDF. If the average extracted text per page is
#     suspiciously low, assume the PDF is scanned and fall back to OCR.
#     """
#     reader = PdfReader(file_path)
#     num_pages = len(reader.pages)
#     text_parts = [page.extract_text() or "" for page in reader.pages]
#     text = "\n".join(text_parts).strip()

#     avg_chars_per_page = len(text) / max(num_pages, 1)
#     if avg_chars_per_page < Config.OCR_TEXT_LENGTH_THRESHOLD:
#         logger.info(
#             "PDF appears scanned (avg %.1f chars/page) - falling back to OCR",
#             avg_chars_per_page,
#         )
#         text = ocr_pdf(file_path)
#     else:
#         logger.info("Extracted %d characters directly from PDF text layer", len(text))
#     return text


# def ocr_pdf(file_path: str) -> str:
#     """
#     Run OCR on every page of a scanned PDF.
#     Requires the Tesseract OCR binary and Poppler to be installed locally
#     (these are system packages, not pip packages).
#     """
#     try:
#         images = convert_from_path(file_path)
#     except Exception as e:
#         logger.error("Failed to convert PDF to images for OCR: %s", e)
#         raise

#     ocr_text = []
#     for i, image in enumerate(images, start=1):
#         logger.info("Running OCR on page %d/%d", i, len(images))
#         ocr_text.append(pytesseract.image_to_string(image))
#     return "\n".join(ocr_text).strip()


# def extract_text_from_image(file_path: str) -> str:
#     """Run OCR directly on an uploaded image file (png/jpg/jpeg)."""
#     logger.info("Running OCR on image file: %s", file_path)
#     image = Image.open(file_path)
#     return pytesseract.image_to_string(image).strip()


# def extract_text_from_docx(file_path: str) -> str:
#     """Extract paragraph text from a Word document."""
#     document = docx.Document(file_path)
#     text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
#     logger.info("Extracted %d characters from DOCX", len(text))
#     return text


# def extract_text_from_csv(file_path: str) -> str:
#     """Convert a CSV file into a row-by-row text representation for retrieval."""
#     df = pd.read_csv(file_path)
#     lines = [f"Columns: {', '.join(df.columns.astype(str))}"]
#     for idx, row in df.iterrows():
#         row_text = ", ".join(f"{col}={row[col]}" for col in df.columns)
#         lines.append(f"Row {idx}: {row_text}")
#     text = "\n".join(lines)
#     logger.info("Converted CSV with %d rows into text (%d characters)", len(df), len(text))
#     return text


# def extract_text_from_excel(file_path: str) -> str:
#     """Convert every sheet of an Excel file into a row-by-row text representation."""
#     sheets = pd.read_excel(file_path, sheet_name=None)
#     all_lines = []
#     for sheet_name, df in sheets.items():
#         all_lines.append(f"--- Sheet: {sheet_name} ---")
#         all_lines.append(f"Columns: {', '.join(df.columns.astype(str))}")
#         for idx, row in df.iterrows():
#             row_text = ", ".join(f"{col}={row[col]}" for col in df.columns)
#             all_lines.append(f"Row {idx}: {row_text}")
#     text = "\n".join(all_lines)
#     logger.info(
#         "Converted Excel file with %d sheet(s) into text (%d characters)",
#         len(sheets),
#         len(text),
#     )
#     return text


# def extract_text_from_txt(file_path: str) -> str:
#     """Read a plain text file."""
#     with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#         text = f.read()
#     logger.info("Read %d characters from TXT file", len(text))
#     return text


# def load_document(file_path: str, file_name: str) -> str:
#     """
#     Detect a document's type from its file extension and route it to the
#     correct text extractor.
#     Supported: pdf (text or scanned), docx, csv, xlsx/xls, txt, png/jpg/jpeg.
#     """
#     ext = os.path.splitext(file_name)[1].lower()
#     logger.info("Loading document '%s' (detected type: %s)", file_name, ext)

#     if ext == ".pdf":
#         return extract_text_from_pdf(file_path)
#     elif ext == ".docx":
#         return extract_text_from_docx(file_path)
#     elif ext == ".csv":
#         return extract_text_from_csv(file_path)
#     elif ext in (".xlsx", ".xls"):
#         return extract_text_from_excel(file_path)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_path)
#     elif ext in (".png", ".jpg", ".jpeg"):
#         return extract_text_from_image(file_path)
#     else:
#         logger.warning("Unsupported file type: %s", ext)
#         raise ValueError(f"Unsupported file type: {ext}")


# # ---------------------------------------------------------------------------
# # Vector store (FAISS)
# # ---------------------------------------------------------------------------
# _embeddings_instance: Optional[HuggingFaceEmbeddings] = None


# def get_embeddings() -> HuggingFaceEmbeddings:
#     """Lazily load and cache the local embedding model (loaded once per process)."""
#     global _embeddings_instance
#     if _embeddings_instance is None:
#         logger.info("Loading embedding model: %s", Config.EMBEDDING_MODEL_NAME)
#         _embeddings_instance = HuggingFaceEmbeddings(model_name=Config.EMBEDDING_MODEL_NAME)
#     return _embeddings_instance


# def build_or_update_vectorstore(
#     text: str, source_name: str, vectorstore: Optional[FAISS] = None
# ) -> FAISS:
#     """
#     Split text into chunks, embed them, and add them to a FAISS vector store.
#     Creates a new store if one doesn't exist yet, otherwise merges into it.
#     """
#     splitter = RecursiveCharacterTextSplitter(
#         chunk_size=Config.CHUNK_SIZE,
#         chunk_overlap=Config.CHUNK_OVERLAP,
#     )
#     chunks = splitter.split_text(text)
#     logger.info("Split '%s' into %d chunks", source_name, len(chunks))

#     documents = [
#         Document(page_content=chunk, metadata={"source": source_name}) for chunk in chunks
#     ]

#     embeddings = get_embeddings()
#     if vectorstore is None:
#         logger.info("Creating new FAISS vector store")
#         vectorstore = FAISS.from_documents(documents, embeddings)
#     else:
#         logger.info("Adding %d chunks to existing FAISS vector store", len(documents))
#         vectorstore.add_documents(documents)

#     return vectorstore


# def process_uploaded_file(
#     file_path: str, file_name: str, vectorstore: Optional[FAISS]
# ) -> FAISS:
#     """
#     End-to-end ingestion of a single uploaded file: extract its text, then
#     chunk/embed/add it to the vector store (creating one if needed).
#     """
#     text = load_document(file_path, file_name)
#     if not text.strip():
#         logger.warning("No text could be extracted from '%s'", file_name)
#         raise ValueError(f"No text could be extracted from '{file_name}'.")
#     return build_or_update_vectorstore(text, file_name, vectorstore)


# def retrieve_context(vectorstore: Optional[FAISS], query: str, k: int = Config.TOP_K) -> List[Document]:
#     """Retrieve the top-k most relevant chunks for a query."""
#     if vectorstore is None:
#         return []
#     logger.info("Retrieving top-%d chunks for query: %s", k, query)
#     results = vectorstore.similarity_search(query, k=k)
#     logger.info("Retrieved %d chunks", len(results))
#     return results


# def save_vectorstore(vectorstore: FAISS, path: str = Config.VECTORSTORE_DIR) -> None:
#     """Persist the FAISS index to disk (optional - session state already holds it in memory)."""
#     vectorstore.save_local(path)
#     logger.info("Vector store saved to '%s'", path)


# def load_vectorstore(path: str = Config.VECTORSTORE_DIR) -> Optional[FAISS]:
#     """Load a previously persisted FAISS index from disk, if it exists."""
#     if not os.path.exists(path):
#         logger.info("No existing vector store found at '%s'", path)
#         return None
#     logger.info("Loading vector store from '%s'", path)
#     return FAISS.load_local(path, get_embeddings(), allow_dangerous_deserialization=True)


# # ---------------------------------------------------------------------------
# # Web search + scraping
# # ---------------------------------------------------------------------------
# def web_search(query: str, max_results: int = Config.WEB_SEARCH_MAX_RESULTS) -> List[Dict[str, str]]:
#     """
#     Search the web via DuckDuckGo (no API key required).
#     Returns a list of dicts with 'title', 'href', and 'body' (snippet) keys.
#     """
#     logger.info("Running web search for: %s", query)
#     try:
#         with DDGS() as ddgs:
#             results = list(ddgs.text(query, max_results=max_results))
#         logger.info("Web search returned %d results", len(results))
#         return results
#     except Exception as e:
#         logger.error("Web search failed: %s", e)
#         return []


# def scrape_url(url: str, timeout: int = Config.SCRAPE_TIMEOUT) -> str:
#     """Fetch a URL and extract its visible text content using BeautifulSoup."""
#     logger.info("Scraping URL: %s", url)
#     try:
#         response = requests.get(url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
#         response.raise_for_status()
#         soup = BeautifulSoup(response.text, "html.parser")

#         # Strip tags that aren't useful as readable content.
#         for tag in soup(["script", "style", "nav", "footer", "header"]):
#             tag.decompose()

#         text = " ".join(soup.stripped_strings)
#         logger.info("Scraped %d characters from %s", len(text), url)
#         return text[: Config.SCRAPE_MAX_CHARS]
#     except requests.exceptions.RequestException as e:
#         logger.warning("Failed to scrape %s: %s", url, e)
#         return ""


# def gather_web_context(query: str) -> str:
#     """
#     Search DuckDuckGo, then scrape the top result pages for fuller content
#     (falling back to the snippet if scraping fails). Returns one combined
#     text block to use as context.
#     """
#     results = web_search(query)
#     if not results:
#         return ""

#     context_parts = []
#     for result in results:
#         snippet = result.get("body", "")
#         url = result.get("href", "")
#         scraped = scrape_url(url) if url else ""
#         combined = scraped if scraped else snippet
#         if combined:
#             context_parts.append(f"Source: {url}\n{combined}")

#     return "\n\n".join(context_parts)


# # ---------------------------------------------------------------------------
# # RAG orchestration
# # ---------------------------------------------------------------------------
# def build_prompt(
#     chat_history: List[Dict[str, str]],
#     doc_context: str,
#     web_context: str,
#     question: str,
# ) -> str:
#     """
#     Assemble the final prompt sent to the LLM: a system instruction, any
#     retrieved document/web context, recent chat history, and the question.
#     """
#     history_text = "\n".join(
#         f"{turn['role'].capitalize()}: {turn['content']}" for turn in chat_history[-6:]
#     )

#     context_blocks = []
#     if doc_context:
#         context_blocks.append(f"Document context:\n{doc_context}")
#     if web_context:
#         context_blocks.append(f"Web search context:\n{web_context}")
#     context_text = "\n\n".join(context_blocks) if context_blocks else "No additional context retrieved."

#     prompt = (
#         "You are a helpful assistant. Answer the user's question using the context "
#         "below when it is relevant. If the context doesn't contain the answer, say so "
#         "honestly instead of making something up.\n\n"
#         f"{context_text}\n\n"
#         f"Conversation so far:\n{history_text}\n\n"
#         f"User: {question}\n"
#         "Assistant:"
#     )
#     return prompt


# def generate_answer(
#     llm: QwenNgrokLLM,
#     chat_history: List[Dict[str, str]],
#     vectorstore: Optional[FAISS],
#     question: str,
#     use_web_search: bool,
# ) -> Tuple[str, List[str]]:
#     """
#     Run the full RAG pipeline for a single question:
#       1. Retrieve relevant document chunks (if a vector store exists)
#       2. Optionally gather web search + scraped context
#       3. Build the final prompt
#       4. Call the LLM
#     Returns (answer, sources_used).
#     """
#     sources: List[str] = []

#     doc_context = ""
#     if vectorstore is not None:
#         chunks = retrieve_context(vectorstore, question)
#         if chunks:
#             doc_context = "\n\n".join(c.page_content for c in chunks)
#             sources.extend(sorted({c.metadata.get("source", "document") for c in chunks}))

#     web_context = ""
#     if use_web_search:
#         web_context = gather_web_context(question)
#         if web_context:
#             sources.append("web search")

#     prompt = build_prompt(chat_history, doc_context, web_context, question)
#     logger.info("Sending prompt to LLM (%d characters)", len(prompt))
#     answer = llm.invoke(prompt)

#     return answer, sources


# # ---------------------------------------------------------------------------
# # Quick standalone self-test (run "python backend.py" to sanity-check the
# # LLM connection without going through the Streamlit UI)
# # ---------------------------------------------------------------------------
# if __name__ == "__main__":
#     logger.info("Running backend self-test...")
#     if check_llm_health():
#         logger.info("LLM endpoint is reachable.")
#     else:
#         logger.warning("LLM endpoint is NOT reachable. Check Config.NGROK_URL.")










# ---------------------- works like a table of contents for the file, making it easy to navigate and understand the structure at a glance ----------------------







# """
# backend.py
# ===========
# Core RAG backend for Groot + LangGraph execution pipeline.

# Sections:
#   1. Config
#   2. Groot System Prompt (F4)
#   3. Qwen LLM wrapper
#   4. Document loading & text extraction
#   5. Semantic chunking with embeddings (F3)
#   6. Tabular ingestion - Excel & CSV
#   7. FAISS vector store
#   8. Web search + scraping
#   9. Retrieval quality & routing
#   10. Follow-up detection (F6)
#   11. Conversation memory (F1)
#   12. Title generation (F5)
#   13. Prompt builder (F7)
#   14. LangGraph pipeline (F2)
#   15. Public generate_answer entry point
# """

# import logging
# import os
# import re
# from dataclasses import dataclass
# from typing import Any, Dict, List, Mapping, Optional, Tuple, TypedDict

# import numpy as np
# import pandas as pd
# import requests
# from bs4 import BeautifulSoup
# from openpyxl import load_workbook
# from PIL import Image
# from pdf2image import convert_from_path
# from pypdf import PdfReader
# import docx  # python-docx
# import pytesseract

# pytesseract.pytesseract.tesseract_cmd = (
#     r"C:\Users\KULDEEP.AMRELIYA\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
# )


# from langchain_core.callbacks.manager import CallbackManagerForLLMRun
# from langchain_core.documents import Document
# from langchain_core.language_models.llms import LLM
# from langchain_community.vectorstores import FAISS
# from langchain_huggingface import HuggingFaceEmbeddings
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langgraph.graph import END, StateGraph

# import db
# import memory_cache

# try:
#     from ddgs import DDGS
# except ImportError:
#     from duckduckgo_search import DDGS


# # ===========================================================================
# # 1. Logging
# # ===========================================================================
# logging.basicConfig(
#     level=logging.INFO,
#     format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
#     handlers=[
#         logging.StreamHandler(),
#         logging.FileHandler("rag_chatbot.log", encoding="utf-8"),
#     ],
# )
# logger = logging.getLogger("rag_backend")


# # ===========================================================================
# # 2. Config
# # ===========================================================================
# class Config:
#     NGROK_URL: str = os.getenv("NGROK_URL", "")
#     APP_API_KEY: str = os.getenv("APP_API_KEY", "")

#     EMBEDDING_MODEL_NAME: str = "sentence-transformers/all-MiniLM-L6-v2"
#     FAISS_BASE_DIR: str = "faiss_indexes"

#     # Chunking
#     CHUNK_SIZE: int = 600
#     CHUNK_OVERLAP: int = 80
#     ROWS_PER_CHUNK: int = 20
#     SEMANTIC_MAX_SENTENCES: int = 8
#     SEMANTIC_MIN_CHARS: int = 120
#     SEMANTIC_SIMILARITY_THRESHOLD: float = 0.55  # tune per corpus; lower = larger chunks

#     # Retrieval
#     TOP_K: int = 3
#     SIMILARITY_SCORE_THRESHOLD: float = 1.0
#     OCR_TEXT_LENGTH_THRESHOLD: int = 20

#     # LLM
#     MAX_NEW_TOKENS: int = 512
#     TEMPERATURE: float = 0.2
#     LLM_REQUEST_TIMEOUT: int = 120

#     # Web search
#     WEB_SEARCH_MAX_RESULTS: int = 3
#     SCRAPE_TIMEOUT: int = 10
#     SCRAPE_MAX_CHARS: int = 2000

#     # Memory
#     SHORT_TERM_WINDOW: int = 6
#     ENABLE_SUMMARY_MEMORY: bool = True
#     SUMMARY_UPDATE_INTERVAL: int = 10

#     # F8: Token budget (~1750 tokens safe for Qwen2.5-3B on T4)
#     MAX_PROMPT_CHARS: int = 7000


# # ===========================================================================
# # 3. Groot System Prompt (F4)
# # ===========================================================================
# SYSTEM_PROMPT = """You are Groot — a friendly, intelligent, and helpful AI assistant.

# Your personality:
# - Warm and conversational, but always professional when the topic demands it.
# - Occasionally witty — a well-placed joke or reference is welcome, but NEVER at the cost of accuracy.
# - You may sometimes reference cricket, Bollywood, Indian culture, or internet memes — but only when it fits naturally. Do NOT force it.
# - When someone asks "Who are you?", introduce yourself as Groot and briefly explain what you can do.

# Your rules:
# 1. Accuracy first. Never sacrifice correctness for humor.
# 2. If you don't know something, say so honestly. Do not hallucinate.
# 3. Use the provided context (documents, web results, conversation history) to answer grounded in facts.
# 4. Keep answers concise unless the user explicitly asks for detail.
# 5. For technical questions, be precise and structured.
# 6. If humor doesn't fit (medical, legal, emotional topics), skip it entirely.

# You have access to:
# - Uploaded documents (PDFs, Word files, spreadsheets, images)
# - Web search results (when documents don't have the answer)
# - Full conversation history
# """


# # ===========================================================================
# # 4. Qwen LLM Wrapper
# # ===========================================================================
# class QwenNgrokLLM(LLM):
#     """LangChain-compatible wrapper for Qwen2.5-3B-Instruct on Kaggle/ngrok."""

#     ngrok_url: str
#     api_key: str
#     max_new_tokens: int = Config.MAX_NEW_TOKENS
#     temperature: float = Config.TEMPERATURE
#     timeout: int = Config.LLM_REQUEST_TIMEOUT

#     @property
#     def _llm_type(self) -> str:
#         return "qwen-ngrok"

#     def _call(
#         self,
#         prompt: str,
#         stop: Optional[List[str]] = None,
#         run_manager: Optional[CallbackManagerForLLMRun] = None,
#         **kwargs: Any,
#     ) -> str:
#         logger.info("Calling Qwen LLM (prompt=%d chars)", len(prompt))
#         try:
#             response = requests.post(
#                 f"{self.ngrok_url}/generate",
#                 json={
#                     "prompt": prompt,
#                     "max_new_tokens": kwargs.get("max_new_tokens", self.max_new_tokens),
#                     "temperature": kwargs.get("temperature", self.temperature),
#                 },
#                 headers={"x-api-key": self.api_key},
#                 timeout=self.timeout,
#             )
#             response.raise_for_status()
#             text = response.json()["response"]
#             logger.info("LLM responded (%d chars)", len(text))
#             return text
#         except requests.exceptions.RequestException as e:
#             logger.error("LLM request failed: %s", e)
#             # Raise instead of returning a fake answer string.
#             # This prevents transient network errors from polluting the
#             # conversation history and memory context.
#             raise RuntimeError(f"LLM unreachable: {e}") from e

#     @property
#     def _identifying_params(self) -> Mapping[str, Any]:
#         return {"ngrok_url": self.ngrok_url, "max_new_tokens": self.max_new_tokens}


# def create_llm() -> QwenNgrokLLM:
#     if not Config.NGROK_URL or not Config.APP_API_KEY:
#         logger.warning("NGROK_URL and/or APP_API_KEY not set — chat will fail until configured in .env.")
#     return QwenNgrokLLM(
#         ngrok_url=Config.NGROK_URL,
#         api_key=Config.APP_API_KEY,
#         max_new_tokens=Config.MAX_NEW_TOKENS,
#         temperature=Config.TEMPERATURE,
#     )


# def check_llm_health() -> bool:
#     if not Config.NGROK_URL:
#         return False
#     try:
#         r = requests.get(f"{Config.NGROK_URL}/health", timeout=10)
#         return r.status_code == 200
#     except requests.exceptions.RequestException as e:
#         logger.warning("LLM health check failed: %s", e)
#         return False


# # ===========================================================================
# # 5. Document Loading & Text Extraction
# # ===========================================================================
# def extract_text_from_pdf(file_path: str) -> str:
#     reader = PdfReader(file_path)
#     num_pages = len(reader.pages)
#     text_parts = [page.extract_text() or "" for page in reader.pages]
#     text = "\n".join(text_parts).strip()
#     avg_chars = len(text) / max(num_pages, 1)
#     if avg_chars < Config.OCR_TEXT_LENGTH_THRESHOLD:
#         logger.info("PDF appears scanned (avg %.1f chars/page) — using OCR", avg_chars)
#         return ocr_pdf(file_path)
#     logger.info("Extracted %d chars from PDF text layer", len(text))
#     return text


# def ocr_pdf(file_path: str) -> str:
#     try:
#         images = convert_from_path(
#             file_path,
#             poppler_path=r"C:\poppler\poppler-26.02.0\Library\bin"
#         )
#     except Exception as e:
#         logger.error("PDF→image conversion failed: %s", e)
#         raise
#     ocr_text = []
#     for i, image in enumerate(images, start=1):
#         logger.info("OCR page %d/%d", i, len(images))
#         ocr_text.append(pytesseract.image_to_string(image))
#     return "\n".join(ocr_text).strip()


# def extract_text_from_image(file_path: str) -> str:
#     logger.info("OCR on image: %s", file_path)
#     return pytesseract.image_to_string(Image.open(file_path)).strip()


# def extract_text_from_docx(file_path: str) -> str:
#     document = docx.Document(file_path)
#     text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
#     logger.info("Extracted %d chars from DOCX", len(text))
#     return text


# def extract_text_from_txt(file_path: str) -> str:
#     with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
#         text = f.read()
#     logger.info("Read %d chars from TXT", len(text))
#     return text


# def load_document(file_path: str, file_name: str) -> str:
#     ext = os.path.splitext(file_name)[1].lower()
#     logger.info("Loading '%s' (type: %s)", file_name, ext)
#     if ext == ".pdf":
#         return extract_text_from_pdf(file_path)
#     elif ext == ".docx":
#         return extract_text_from_docx(file_path)
#     elif ext == ".txt":
#         return extract_text_from_txt(file_path)
#     elif ext in (".png", ".jpg", ".jpeg"):
#         return extract_text_from_image(file_path)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")


# # ===========================================================================
# # 6. Semantic Chunking with Embeddings for PDF / DOCX / TXT (F3)
# # ===========================================================================
# # Strategy: 
# #   1. Regex sentence splitter + section header detection (structural signals)
# #   2. Embedding-similarity grouping within each section (topic-aware splits)
# # No dependency bloat — reuses the same all-MiniLM-L6-v2 model already
# # loaded for FAISS. Chunking happens only at ingest time, not query time.

# _SECTION_HEADER_RE = re.compile(
#     r"^(?:"
#     r"\d+(?:\.\d+)*[\s\.\)]+[A-Z]"   # "1. Title" or "2.1 Section"
#     r"|[A-Z][A-Z\s]{4,}$"            # ALL CAPS line (≥5 chars)
#     r"|.{3,60}:$"                     # "Some label:" at line end
#     r")",
#     re.MULTILINE,
# )

# _SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")


# def _split_sentences(text: str) -> List[str]:
#     sentences = _SENTENCE_SPLIT_RE.split(text)
#     result = []
#     for s in sentences:
#         for line in s.split("\n"):
#             line = line.strip()
#             if line:
#                 result.append(line)
#     return result


# def _is_section_header(line: str) -> bool:
#     return bool(_SECTION_HEADER_RE.match(line.strip()))


# def _cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
#     """Cosine similarity between two vectors, safe for zero vectors."""
#     denom = np.linalg.norm(a) * np.linalg.norm(b)
#     if denom == 0:
#         return 0.0
#     return float(np.dot(a, b) / denom)


# def _semantic_group_sentences(
#     sentences: List[str],
#     similarity_threshold: float = Config.SEMANTIC_SIMILARITY_THRESHOLD,
# ) -> List[List[str]]:
#     """
#     Group sentences into semantically coherent runs using embedding
#     similarity between consecutive sentences. A new group starts when:
#       1. Similarity drops below threshold (topic shift)
#       2. Size/char caps are hit (prevent pathologically large chunks)
#     """
#     if not sentences:
#         return []
#     if len(sentences) == 1:
#         return [sentences]

#     embeddings = get_embeddings().embed_documents(sentences)
#     embeddings = [np.array(e) for e in embeddings]

#     groups: List[List[str]] = [[sentences[0]]]
#     group_chars = len(sentences[0])

#     for i in range(1, len(sentences)):
#         sim = _cosine_sim(embeddings[i - 1], embeddings[i])
#         sentence = sentences[i]
#         exceeds_caps = (
#             len(groups[-1]) >= Config.SEMANTIC_MAX_SENTENCES
#             or group_chars + len(sentence) > Config.CHUNK_SIZE
#         )
#         if sim < similarity_threshold or exceeds_caps:
#             groups.append([sentence])
#             group_chars = len(sentence)
#         else:
#             groups[-1].append(sentence)
#             group_chars += len(sentence)

#     return groups


# def semantic_chunk_text(text: str, source_name: str) -> List[Document]:
#     """
#     Split text into semantically coherent chunks using:
#       1. Regex section-header detection (hard boundaries).
#       2. Embedding-similarity grouping within each section (soft boundaries).

#     Returns LangChain Documents with source and section metadata.
#     """
#     sentences = _split_sentences(text)
#     if not sentences:
#         return []

#     chunks: List[Document] = []
#     current_section: str = "Introduction"
#     section_sentences: List[str] = []

#     def flush_section(section: str) -> None:
#         nonlocal section_sentences
#         if not section_sentences:
#             return
#         for group in _semantic_group_sentences(section_sentences):
#             content = " ".join(group)
#             if len(content) >= Config.SEMANTIC_MIN_CHARS:
#                 chunks.append(Document(
#                     page_content=content,
#                     metadata={"source": source_name, "section": section},
#                 ))
#             elif chunks:
#                 # Merge tiny trailing chunk into previous
#                 chunks[-1] = Document(
#                     page_content=chunks[-1].page_content + " " + content,
#                     metadata=chunks[-1].metadata,
#                 )
#             else:
#                 # Tiny chunk with no prior chunks — add it anyway
#                 chunks.append(Document(
#                     page_content=content,
#                     metadata={"source": source_name, "section": section},
#                 ))
#         section_sentences = []

#     for sentence in sentences:
#         if _is_section_header(sentence):
#             flush_section(current_section)
#             current_section = sentence.strip()
#             continue
#         section_sentences.append(sentence)

#     flush_section(current_section)

#     logger.info("Semantic chunking: '%s' → %d chunk(s)", source_name, len(chunks))
#     return chunks


# # ===========================================================================
# # 7. Tabular Ingestion — Excel & CSV
# # ===========================================================================
# def extract_rows_from_excel(file_path: str) -> List[Dict[str, Any]]:
#     logger.info("Opening Excel workbook: %s", file_path)
#     workbook = load_workbook(file_path, read_only=True, data_only=True)
#     rows_out: List[Dict[str, Any]] = []

#     for sheet in workbook.worksheets:
#         sheet_name = sheet.title
#         rows_iter = sheet.iter_rows(values_only=True)
#         try:
#             header = next(rows_iter)
#         except StopIteration:
#             logger.warning("Sheet '%s' is empty — skipping", sheet_name)
#             continue

#         headers = [
#             str(h).strip() if h is not None else f"col_{i}"
#             for i, h in enumerate(header)
#         ]
#         for row_idx, row in enumerate(rows_iter, start=2):
#             if row is None or all(cell is None for cell in row):
#                 continue
#             pairs = [
#                 f"{col} is {val}"
#                 for col, val in zip(headers, row)
#                 if val is not None and str(val).strip()
#             ]
#             if not pairs:
#                 continue
#             rows_out.append({
#                 "sheet": sheet_name,
#                 "row_index": row_idx,
#                 "text": f"In sheet '{sheet_name}', row {row_idx}: " + "; ".join(pairs) + ".",
#             })

#     workbook.close()
#     logger.info("Excel: %d row(s) extracted", len(rows_out))
#     return rows_out


# def extract_rows_from_csv(file_path: str, read_chunksize: int = 5000) -> List[Dict[str, Any]]:
#     rows_out: List[Dict[str, Any]] = []
#     for batch_idx, df_chunk in enumerate(pd.read_csv(file_path, chunksize=read_chunksize), start=1):
#         columns = [str(c).strip() for c in df_chunk.columns]
#         for idx, row in df_chunk.iterrows():
#             pairs = [
#                 f"{col} is {row[col]}"
#                 for col in columns
#                 if pd.notna(row[col]) and str(row[col]).strip()
#             ]
#             if not pairs:
#                 continue
#             rows_out.append({
#                 "sheet": "csv",
#                 "row_index": int(idx),
#                 "text": f"Row {idx}: " + "; ".join(pairs) + ".",
#             })
#         logger.info("CSV batch %d: %d row(s) total", batch_idx, len(rows_out))
#     logger.info("CSV: %d row(s) extracted", len(rows_out))
#     return rows_out


# def chunk_tabular_rows(
#     rows: List[Dict[str, Any]],
#     source_name: str,
#     rows_per_chunk: int = Config.ROWS_PER_CHUNK,
# ) -> List[Document]:
#     documents: List[Document] = []
#     for i in range(0, len(rows), rows_per_chunk):
#         batch = rows[i: i + rows_per_chunk]
#         documents.append(Document(
#             page_content="\n".join(r["text"] for r in batch),
#             metadata={
#                 "source": source_name,
#                 "sheets": ", ".join(sorted({r["sheet"] for r in batch})),
#                 "row_start": batch[0]["row_index"],
#                 "row_end": batch[-1]["row_index"],
#             },
#         ))
#     logger.info("Tabular: %d row(s) → %d chunk(s) for '%s'", len(rows), len(documents), source_name)
#     return documents


# def process_uploaded_file(
#     file_path: str,
#     file_name: str,
#     vectorstore: Optional[FAISS],
# ) -> Tuple[FAISS, int]:
#     ext = os.path.splitext(file_name)[1].lower()

#     if ext in (".xlsx", ".xls"):
#         rows = extract_rows_from_excel(file_path)
#         if not rows:
#             raise ValueError(f"No data rows extracted from '{file_name}'.")
#         documents = chunk_tabular_rows(rows, file_name)

#     elif ext == ".csv":
#         rows = extract_rows_from_csv(file_path)
#         if not rows:
#             raise ValueError(f"No data rows extracted from '{file_name}'.")
#         documents = chunk_tabular_rows(rows, file_name)

#     else:
#         text = load_document(file_path, file_name)
#         if not text.strip():
#             raise ValueError(f"No text extracted from '{file_name}'.")
#         documents = semantic_chunk_text(text, file_name)
#         if not documents:
#             logger.warning("Semantic chunker returned 0 chunks — falling back to RecursiveCharacterTextSplitter")
#             splitter = RecursiveCharacterTextSplitter(
#                 chunk_size=Config.CHUNK_SIZE,
#                 chunk_overlap=Config.CHUNK_OVERLAP,
#             )
#             documents = [
#                 Document(page_content=c, metadata={"source": file_name})
#                 for c in splitter.split_text(text)
#             ]

#     vectorstore = add_documents_to_vectorstore(documents, vectorstore)
#     return vectorstore, len(documents)


# # ===========================================================================
# # 8. FAISS Vector Store
# # ===========================================================================
# _embeddings_instance: Optional[HuggingFaceEmbeddings] = None


# def get_embeddings() -> HuggingFaceEmbeddings:
#     global _embeddings_instance
#     if _embeddings_instance is None:
#         logger.info("Loading embedding model: %s", Config.EMBEDDING_MODEL_NAME)
#         _embeddings_instance = HuggingFaceEmbeddings(model_name=Config.EMBEDDING_MODEL_NAME)
#     return _embeddings_instance


# def get_vectorstore_path(conversation_id: str) -> str:
#     return os.path.join(Config.FAISS_BASE_DIR, str(conversation_id))


# def add_documents_to_vectorstore(
#     documents: List[Document],
#     vectorstore: Optional[FAISS],
# ) -> FAISS:
#     embeddings = get_embeddings()
#     if vectorstore is None:
#         logger.info("Creating new FAISS index (%d chunks)", len(documents))
#         return FAISS.from_documents(documents, embeddings)
#     logger.info("Adding %d chunks to FAISS index", len(documents))
#     vectorstore.add_documents(documents)
#     return vectorstore


# def save_vectorstore(vectorstore: FAISS, path: str) -> None:
#     parent = os.path.dirname(path)
#     if parent:
#         os.makedirs(parent, exist_ok=True)
#     vectorstore.save_local(path)
#     logger.info("Vector store saved: '%s'", path)


# def load_vectorstore(path: str) -> Optional[FAISS]:
#     if not os.path.exists(path):
#         return None
#     logger.info("Loading vector store: '%s'", path)
#     return FAISS.load_local(path, get_embeddings(), allow_dangerous_deserialization=True)


# def get_vectorstore_status(
#     vectorstore: Optional[FAISS],
#     faiss_path: str,
#     conversation_id: str,
# ) -> Dict[str, Any]:
#     chunk_count = vectorstore.index.ntotal if vectorstore is not None else 0
#     try:
#         document_count = len(db.list_documents(conversation_id))
#     except Exception as e:
#         logger.error("Could not fetch document count: %s", e)
#         document_count = 0
#     return {
#         "loaded": vectorstore is not None,
#         "embedding_model": Config.EMBEDDING_MODEL_NAME,
#         "document_count": document_count,
#         "chunk_count": chunk_count,
#         "faiss_path": faiss_path,
#     }


# # ===========================================================================
# # 9. Web Search + Scraping
# # ===========================================================================
# def web_search(query: str, max_results: int = Config.WEB_SEARCH_MAX_RESULTS) -> List[Dict[str, str]]:
#     logger.info("Web search: %s", query)
#     try:
#         with DDGS() as ddgs:
#             results = list(ddgs.text(query, max_results=max_results))
#         logger.info("Web search: %d result(s)", len(results))
#         return results
#     except Exception as e:
#         logger.error("Web search failed: %s", e)
#         return []


# def scrape_url(url: str, timeout: int = Config.SCRAPE_TIMEOUT) -> str:
#     logger.info("Scraping: %s", url)
#     try:
#         response = requests.get(url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"})
#         response.raise_for_status()
#         soup = BeautifulSoup(response.text, "html.parser")
#         for tag in soup(["script", "style", "nav", "footer", "header"]):
#             tag.decompose()
#         return " ".join(soup.stripped_strings)[: Config.SCRAPE_MAX_CHARS]
#     except requests.exceptions.RequestException as e:
#         logger.warning("Scrape failed for %s: %s", url, e)
#         return ""


# def gather_web_context(query: str) -> str:
#     results = web_search(query)
#     if not results:
#         return ""
#     parts = []
#     for result in results:
#         url = result.get("href", "")
#         content = scrape_url(url) if url else result.get("body", "")
#         if content:
#             parts.append(f"Source: {url}\n{content}")
#     return "\n\n".join(parts)


# # ===========================================================================
# # 10. Retrieval Quality & Routing
# # ===========================================================================
# def assess_retrieval_quality(
#     vectorstore: Optional[FAISS],
#     query: str,
#     k: int = Config.TOP_K,
# ) -> Tuple[List[Document], bool]:
#     if vectorstore is None:
#         return [], False
#     try:
#         results = vectorstore.similarity_search_with_score(query, k=k)
#     except Exception as e:
#         logger.error("Retrieval failed: %s", e)
#         return [], False
#     if not results:
#         return [], False
#     docs = [doc for doc, _ in results]
#     best_score = min(score for _, score in results)
#     is_strong = best_score <= Config.SIMILARITY_SCORE_THRESHOLD
#     logger.info("Retrieval: best_score=%.4f strong=%s", best_score, is_strong)
#     return docs, is_strong


# def should_use_web_search(vectorstore: Optional[FAISS], retrieval_strong: bool) -> bool:
#     return vectorstore is None or not retrieval_strong


# # ===========================================================================
# # 11. Follow-up Detection (F6)
# # ===========================================================================
# _FOLLOWUP_TRIGGERS = re.compile(
#     r"^("
#     r"(what|how|why|when|where|who|can|could|should|would|does|is|are|do|did)\s+(it|this|that|they|he|she|we)\b"
#     r"|tell\s+me\s+more"
#     r"|explain\s+(that|this|it|more)"
#     r"|what\s+about"
#     r"|and\s+(what|how|why|can|does)"
#     r"|more\s+(details?|info|examples?)"
#     r"|elaborate|go\s+on|continue"
#     r"|so\s+(it|this|that|they)"
#     r")",
#     re.IGNORECASE,
# )


# def is_follow_up(question: str) -> bool:
#     """Detect if a question is a follow-up to prior context.
    
#     Only uses explicit trigger patterns. Removed the blanket "≤6 words + no caps"
#     heuristic which falsely flagged casual statements like "hey i am jack" as
#     follow-ups, causing unnecessary web searches.
#     """
#     q = question.strip()
#     if _FOLLOWUP_TRIGGERS.match(q):
#         return True
#     return False


# def expand_query_with_context(question: str, recent_messages: list) -> str:
#     """Expand a follow-up query using the LAST PRIOR USER MESSAGE.
    
#     NOT the assistant's last reply (to avoid feedback loops where the bot's
#     own apology/error string gets re-injected as if it were context).
    
#     Args:
#         question: the current user question
#         recent_messages: list of message objects from db/cache with .role/.content
#                         (already includes the current user question as the last item)
    
#     Returns: expanded query string
#     """
#     if not recent_messages or len(recent_messages) < 2:
#         return question
    
#     # Skip the last message (current question) and find the prior user message
#     for m in reversed(recent_messages[:-1]):
#         if m.get("role") == "user" if isinstance(m, dict) else m.role == "user":
#             prior_user_content = m.get("content") if isinstance(m, dict) else m.content
#             expanded = f"{prior_user_content[:200]} {question}"
#             logger.info("Follow-up expanded query: %s", expanded[:120])
#             return expanded
    
#     return question


# # ===========================================================================
# # 12. Conversation Memory (F1) — now with Redis cache
# # ===========================================================================
# @dataclass
# class MemoryMessage:
#     """Unified message format from either Redis (dict) or Postgres (ORM)."""
#     role: str
#     content: str


# def build_memory_context(conversation_id: str) -> Tuple[str, str, list]:
#     """
#     Returns: (short_term_text, long_term_summary, raw_messages)

#     Tries Redis first for the short-term window (cache hit = no Postgres
#     round-trip). Falls back to Postgres on cache miss or Redis being down.

#     raw_messages is normalized to MemoryMessage objects so downstream code
#     (is_follow_up, expand_query_with_context) works unchanged regardless of source.

#     F1 improvements:
#     - Role labels: "User" / "Groot" (not "assistant")
#     - Turns are numbered for sequence clarity
#     - Per-message truncation to 600 chars to control prompt size
#     - raw_messages returned for follow-up detection
#     - Redis short-term cache layer for latency optimization
#     """
#     cached = memory_cache.get_recent_messages(conversation_id)
    
#     if cached is not None:
#         # Redis cache hit
#         messages = [MemoryMessage(role=m["role"], content=m["content"]) for m in cached]
#     else:
#         # Redis miss or unavailable — fall back to Postgres
#         db_messages = db.get_messages(conversation_id, limit=Config.SHORT_TERM_WINDOW)
#         messages = [MemoryMessage(role=m.role, content=m.content) for m in db_messages]

#     lines = []
#     for i, m in enumerate(messages, start=1):
#         role_label = "User" if m.role == "user" else "Groot"
#         content = m.content[:600] if len(m.content) > 600 else m.content
#         lines.append(f"[Turn {i}] {role_label}: {content}")
#     short_term = "\n".join(lines)

#     convo = db.get_conversation(conversation_id)
#     long_term_summary = (convo.summary or "") if convo else ""

#     return short_term, long_term_summary, messages


# def maybe_update_long_term_summary(llm: QwenNgrokLLM, conversation_id: str) -> None:
#     if not Config.ENABLE_SUMMARY_MEMORY:
#         return
#     total = db.count_messages(conversation_id)
#     if total == 0 or total % Config.SUMMARY_UPDATE_INTERVAL != 0:
#         return
#     convo = db.get_conversation(conversation_id)
#     if convo is None:
#         return

#     messages = db.get_messages(conversation_id)
#     history_text = "\n".join(
#         f"{'User' if m.role == 'user' else 'Groot'}: {m.content}"
#         for m in messages
#     )
#     summary_prompt = (
#         "Summarize the following conversation in under 120 words. "
#         "Preserve: key topics discussed, important facts, user preferences, "
#         "and any open questions. Write ONLY the summary, no preamble.\n\n"
#         f"Previous summary:\n{convo.summary or '(none yet)'}\n\n"
#         f"Conversation:\n{history_text}\n\nSummary:"
#     )
#     try:
#         summary = llm.invoke(summary_prompt)
#         db.update_conversation_summary(conversation_id, summary.strip())
#         logger.info("Updated long-term summary for %s (%d messages)", conversation_id, total)
#     except Exception as e:
#         logger.error("Summary update failed for %s: %s", conversation_id, e)


# # ===========================================================================
# # 13. Title Generation (F5)
# # ===========================================================================
# def generate_conversation_title(
#     llm: QwenNgrokLLM,
#     user_question: str,
#     assistant_answer: str,
# ) -> str:
#     """
#     LLM-generated title after first exchange. Runs as a background task
#     in backend_server.py so it never adds latency to the chat response.
#     Falls back to word-truncation on any failure.
#     """
#     prompt = (
#         "Generate a short conversation title (3–5 words, no punctuation) "
#         "that captures the main topic of this exchange.\n\n"
#         f"User: {user_question[:300]}\n"
#         f"Assistant: {assistant_answer[:300]}\n\n"
#         "Title (3-5 words only):"
#     )
#     try:
#         title = llm.invoke(prompt, max_new_tokens=20, temperature=0.3).strip()
#         title = title.split("\n")[0].strip('"\'').strip()
#         if title and 3 <= len(title) <= 80:
#             return title
#     except Exception as e:
#         logger.warning("Title generation failed: %s", e)

#     # Fallback
#     words = user_question.strip().split()
#     title = " ".join(words[:6])
#     if len(words) > 6:
#         title += "..."
#     return title[:80]


# def generate_conversation_title_simple(first_message: str) -> str:
#     """Instant fallback title — no LLM call. Used at conversation creation time."""
#     words = first_message.strip().split()
#     if not words:
#         return "New Chat"
#     title = " ".join(words[:8])
#     if len(words) > 8:
#         title += "..."
#     return title[:100]


# # ===========================================================================
# # 14. Prompt Builder (F7)
# # ===========================================================================
# def _trim_to_budget(text: str, max_chars: int) -> str:
#     if len(text) <= max_chars:
#         return text
#     return text[:max_chars] + "\n[... truncated ...]"


# def build_prompt(
#     short_term_history: str,
#     long_term_summary: str,
#     doc_context: str,
#     web_context: str,
#     question: str,
# ) -> str:
#     summary_block = (
#         "### Conversation Summary (earlier context)\n"
#         + _trim_to_budget(long_term_summary, 600)
#         if long_term_summary else ""
#     )
#     memory_block = (
#         "### Recent Conversation\n"
#         + _trim_to_budget(short_term_history, 1200)
#         if short_term_history else ""
#     )
#     doc_block = (
#         "### Document Context\n"
#         + _trim_to_budget(doc_context, 2000)
#         if doc_context else ""
#     )
#     web_block = (
#         "### Web Search Results\n"
#         + _trim_to_budget(web_context, 1200)
#         if web_context else ""
#     )

#     context_parts = [b for b in [summary_block, memory_block, doc_block, web_block] if b]
#     context_section = "\n\n".join(context_parts) if context_parts else "No additional context available."

#     prompt = (
#         f"{SYSTEM_PROMPT}\n\n"
#         f"{context_section}\n\n"
#         f"### Current Question\n"
#         f"User: {question[:300]}\n"
#         f"Groot:"
#     )
#     logger.info("Prompt built: %d chars", len(prompt))
#     logger.info("=" * 80)
#     logger.info(prompt)
#     logger.info("=" * 80)
    
#     return prompt


# # ===========================================================================
# # 15. LangGraph Pipeline (F2)
# # ===========================================================================

# # ---------------------------------------------------------------------------
# # Graph State
# # ---------------------------------------------------------------------------
# class GraphState(TypedDict):
#     question: str
#     conversation_id: str
#     vectorstore: Optional[FAISS]
#     llm: Any

#     intent: str
#     retrieval_query: str
#     is_follow_up: bool

#     short_term: str
#     long_term_summary: str
#     raw_messages: list

#     doc_context: str
#     retrieval_strong: bool
#     web_context: str

#     answer: str
#     sources: List[str]


# # ---------------------------------------------------------------------------
# # Intent classification — rule-based (no extra LLM call, F8)
# # ---------------------------------------------------------------------------
# _GENERAL_CHAT_RE = re.compile(
#     r"^(hi|hello|hey|howdy)\b"
#     r"|\b(what('s| is) your name|who are you|introduce yourself)\b"
#     r"|^(thanks?|thank you|bye|goodbye)\b"
#     r"|^good (morning|afternoon|evening)\b"
#     r"|\bhow are you\b"
#     r"|\b(my name is|i am|i'm|call me)\s+\w+"
#     r"|what can you do|help me",
#     re.IGNORECASE,
# )

# _WEB_SEARCH_SIGNALS = re.compile(
#     r"\b(latest|current|recent|today|news|2024|2025|live|real.?time|stock price"
#     r"|weather|score|winner|just released|new version|update)\b",
#     re.IGNORECASE,
# )

# _DOCUMENT_SIGNALS = re.compile(
#     r"\b(document|file|uploaded|pdf|report|according to|in the|the document"
#     r"|spreadsheet|table|row|column|sheet|attachment)\b",
#     re.IGNORECASE,
# )


# def _classify_intent(question: str, vectorstore: Optional[FAISS], follow_up: bool) -> str:
#     q = question.strip()
#     if _GENERAL_CHAT_RE.search(q):
#         return "general_chat"
#     if follow_up:
#         return "follow_up_query"
#     if vectorstore is not None and _DOCUMENT_SIGNALS.search(q):
#         return "document_query"
#     if _WEB_SEARCH_SIGNALS.search(q):
#         return "web_search_query"
#     if vectorstore is not None:
#         return "document_query"
#     return "technical_question"


# # ---------------------------------------------------------------------------
# # Nodes
# # ---------------------------------------------------------------------------
# def _intent_classifier_node(state: GraphState) -> GraphState:
#     question = state["question"]
#     follow_up = is_follow_up(question)
#     state["is_follow_up"] = follow_up
#     state["intent"] = _classify_intent(question, state.get("vectorstore"), follow_up)
#     state["retrieval_query"] = question
#     state.setdefault("short_term", "")
#     state.setdefault("long_term_summary", "")
#     state.setdefault("raw_messages", [])
#     state.setdefault("doc_context", "")
#     state.setdefault("retrieval_strong", False)
#     state.setdefault("web_context", "")
#     state.setdefault("sources", [])
#     logger.info("Intent: %s | follow_up: %s", state["intent"], follow_up)
#     return state


# def _memory_node(state: GraphState) -> GraphState:
#     try:
#         short_term, long_term_summary, raw_messages = build_memory_context(state["conversation_id"])
#         state["short_term"] = short_term
#         state["long_term_summary"] = long_term_summary
#         state["raw_messages"] = raw_messages
#     except Exception as e:
#         logger.error("Memory node failed: %s", e)
#         state["short_term"] = ""
#         state["long_term_summary"] = ""
#         state["raw_messages"] = []

#     if state["is_follow_up"] and state["raw_messages"]:
#         state["retrieval_query"] = expand_query_with_context(
#             state["question"], state["raw_messages"]
#         )
#     return state


# def _memory_router(state: GraphState) -> str:
#     intent = state["intent"]
#     if intent == "general_chat":
#         return "general_chat_node"
#     if intent == "web_search_query":
#         return "web_search_node"
#     return "document_retrieval_node"


# def _document_retrieval_node(state: GraphState) -> GraphState:
#     vectorstore = state.get("vectorstore")
#     sources = state.get("sources", [])
#     if vectorstore is None:
#         state["doc_context"] = ""
#         state["retrieval_strong"] = False
#         state["sources"] = sources
#         return state
#     try:
#         docs, retrieval_strong = assess_retrieval_quality(vectorstore, state["retrieval_query"])
#         if docs:
#             state["doc_context"] = "\n\n".join(d.page_content for d in docs)
#             sources.extend(sorted({d.metadata.get("source", "document") for d in docs}))
#         else:
#             state["doc_context"] = ""
#         state["retrieval_strong"] = retrieval_strong
#         state["sources"] = sources
#         logger.info("Doc retrieval: %d chunks, strong=%s", len(docs), retrieval_strong)
#     except Exception as e:
#         logger.error("Document retrieval node failed: %s", e)
#         state["doc_context"] = ""
#         state["retrieval_strong"] = False
#         state["sources"] = sources
#     return state


# def _retrieval_router(state: GraphState) -> str:
#     if should_use_web_search(state.get("vectorstore"), state["retrieval_strong"]):
#         return "web_search_node"
#     return "response_node"


# def _web_search_node(state: GraphState) -> GraphState:
#     sources = state.get("sources", [])
#     try:
#         web_context = gather_web_context(state["retrieval_query"])
#         if web_context:
#             state["web_context"] = web_context
#             sources.append("web search")
#             logger.info("Web search: %d chars", len(web_context))
#         else:
#             state["web_context"] = ""
#     except Exception as e:
#         logger.error("Web search node failed: %s", e)
#         state["web_context"] = ""
#     state["sources"] = sources
#     return state


# def _general_chat_node(state: GraphState) -> GraphState:
#     # No retrieval — memory context is sufficient for greetings/identity
#     logger.info("General chat node: skipping retrieval")
#     return state


# def _response_node(state: GraphState) -> GraphState:
#     prompt = build_prompt(
#         short_term_history=state["short_term"],
#         long_term_summary=state["long_term_summary"],
#         doc_context=state["doc_context"],
#         web_context=state["web_context"],
#         question=state["question"],
#     )
#     try:
#         state["answer"] = state["llm"].invoke(prompt)
#     except Exception as e:
#         logger.error("LLM call failed in response_node: %s", e)
#         raise
#     logger.info("Response node: answer=%d chars", len(state["answer"]))
#     return state


# # ---------------------------------------------------------------------------
# # Graph construction — compiled once at import time
# # ---------------------------------------------------------------------------
# def _build_graph() -> Any:
#     workflow = StateGraph(GraphState)

#     workflow.add_node("intent_classifier_node", _intent_classifier_node)
#     workflow.add_node("memory_node", _memory_node)
#     workflow.add_node("document_retrieval_node", _document_retrieval_node)
#     workflow.add_node("web_search_node", _web_search_node)
#     workflow.add_node("general_chat_node", _general_chat_node)
#     workflow.add_node("response_node", _response_node)

#     workflow.set_entry_point("intent_classifier_node")
#     workflow.add_edge("intent_classifier_node", "memory_node")

#     workflow.add_conditional_edges(
#         "memory_node",
#         _memory_router,
#         {
#             "document_retrieval_node": "document_retrieval_node",
#             "web_search_node": "web_search_node",
#             "general_chat_node": "general_chat_node",
#         },
#     )
#     workflow.add_conditional_edges(
#         "document_retrieval_node",
#         _retrieval_router,
#         {
#             "web_search_node": "web_search_node",
#             "response_node": "response_node",
#         },
#     )
#     workflow.add_edge("web_search_node", "response_node")
#     workflow.add_edge("general_chat_node", "response_node")
#     workflow.add_edge("response_node", END)

#     return workflow.compile()


# _graph = None


# def _get_graph() -> Any:
#     global _graph
#     if _graph is None:
#         _graph = _build_graph()
#     return _graph


# # ===========================================================================
# # 16. Public Entry Point
# # ===========================================================================
# def generate_answer(
#     llm: QwenNgrokLLM,
#     conversation_id: str,
#     vectorstore: Optional[FAISS],
#     question: str,
# ) -> Tuple[str, List[str]]:
#     """
#     Main entry point called by backend_server.py.
#     Runs the LangGraph pipeline; falls back to direct pipeline on graph failure.
#     """
#     initial_state: GraphState = {
#         "question": question,
#         "conversation_id": conversation_id,
#         "vectorstore": vectorstore,
#         "llm": llm,
#         "intent": "",
#         "retrieval_query": question,
#         "is_follow_up": False,
#         "short_term": "",
#         "long_term_summary": "",
#         "raw_messages": [],
#         "doc_context": "",
#         "retrieval_strong": False,
#         "web_context": "",
#         "answer": "",
#         "sources": [],
#     }

#     try:
#         final_state = _get_graph().invoke(initial_state)
#         return final_state["answer"], final_state["sources"]
#     except Exception as e:
#         logger.error("Graph execution failed, running fallback pipeline: %s", e)
#         return _fallback_pipeline(llm, conversation_id, vectorstore, question)


# def _fallback_pipeline(
#     llm: QwenNgrokLLM,
#     conversation_id: str,
#     vectorstore: Optional[FAISS],
#     question: str,
# ) -> Tuple[str, List[str]]:
#     """Direct linear pipeline — used only if LangGraph fails."""
#     sources: List[str] = []
#     try:
#         short_term, long_term_summary, raw_messages = build_memory_context(conversation_id)
#     except Exception:
#         short_term, long_term_summary, raw_messages = "", "", []

#     retrieval_query = expand_query_with_context(question, raw_messages) if is_follow_up(question) else question

#     doc_context = ""
#     retrieval_strong = False
#     try:
#         docs, retrieval_strong = assess_retrieval_quality(vectorstore, retrieval_query)
#         if docs:
#             doc_context = "\n\n".join(d.page_content for d in docs)
#             sources.extend(sorted({d.metadata.get("source", "document") for d in docs}))
#     except Exception:
#         pass

#     web_context = ""
#     if should_use_web_search(vectorstore, retrieval_strong):
#         try:
#             web_context = gather_web_context(retrieval_query)
#             if web_context:
#                 sources.append("web search")
#         except Exception:
#             pass

#     prompt = build_prompt(short_term, long_term_summary, doc_context, web_context, question)
#     try:
#         answer = llm.invoke(prompt)
#     except Exception as e:
#         logger.error("Fallback pipeline LLM call failed: %s", e)
#         raise
#     return answer, sources


# # ===========================================================================
# # Standalone self-test
# # ===========================================================================
# if __name__ == "__main__":
#     logger.info("Backend self-test...")
#     if check_llm_health():
#         logger.info("LLM endpoint reachable.")
#     else:
#         logger.warning("LLM endpoint NOT reachable. Check NGROK_URL in .env.")



















"""
backend.py (v4)
================
Core RAG backend for Groot + LangGraph execution pipeline.

Changes from v3:
  - MEMORY (Section 13): Complete rewrite.
      * _filter_relevant_turns REMOVED — was broken (referenced undefined
        Config constants RELEVANT_TURNS_TOP_K and ALWAYS_INCLUDE_LAST_PAIRS,
        causing AttributeError at import time).
      * SHORT_TERM_WINDOW removed from Config — replaced by SHORT_TERM_CHAR_BUDGET.
      * New _select_pairs_within_budget: packs newest Q&A pairs into a
        1200-char budget (newest→oldest), returned in chronological order.
        Dynamic count — 2-3 pairs on typical exchanges, more on short ones.
      * maybe_update_long_term_summary: now INCREMENTAL — only the last
        SUMMARY_UPDATE_INTERVAL new turns are summarized and merged with
        the prior summary. Never re-summarizes the full history.
      * build_memory_context: fetches up to 20 messages (enough for 10 pairs),
        budget selector trims to fit. No relevance filtering.

  - BUG FIXES:
      * _NO_WEB_SEARCH_INTENTS: removed stale 'follow_up_query' (no longer
        exists), added 'follow_up_chat', 'follow_up_doc', 'knowledge_query'.
      * _memory_router: added explicit 'follow_up_chat' → general_chat_node
        routing (was falling into document_retrieval_node with None vectorstore).
      * Config.SHORT_TERM_CHAR_BUDGET and Config.SUMMARY_CHAR_BUDGET added
        as explicit constants replacing removed SHORT_TERM_WINDOW.

Preserved from v3 (unchanged):
  - Strict 4-tier routing cascade
  - follow_up_doc / follow_up_chat / knowledge_query intents
  - Semantic chunking for ALL document types
  - Tabular analytics: SQL (sqlite3) + math (pandas)
  - FAISS vector store, web search, scraping
  - SIMILARITY_SCORE_THRESHOLD = 0.85

Sections:
  1.  Logging
  2.  Config
  3.  Groot System Prompt (F4)
  4.  Qwen LLM wrapper
  5.  Document loading & text extraction
  6.  Semantic chunking with embeddings (F3)
  7.  Tabular ingestion — Excel & CSV
  8.  Tabular analytics — SQL & math execution
  9.  FAISS vector store
  10. Web search + scraping
  11. Retrieval quality & routing
  12. Follow-up detection (F6)
  13. Conversation memory (F1) — budget-aware pair selection + incremental summary
  14. Title generation (F5)
  15. Prompt builder (F7)
  16. LangGraph pipeline (F2)
  17. Public generate_answer entry point
"""

import logging
import os
import re
import sqlite3
from dataclasses import dataclass
from typing import Any, Dict, List, Mapping, Optional, Tuple, TypedDict

import numpy as np
import pandas as pd
import requests
from bs4 import BeautifulSoup
from openpyxl import load_workbook
from PIL import Image
from pdf2image import convert_from_path
from pypdf import PdfReader
import docx  # python-docx
import pytesseract

pytesseract.pytesseract.tesseract_cmd = (
    r"C:\Users\KULDEEP.AMRELIYA\AppData\Local\Programs\Tesseract-OCR\tesseract.exe"
)

from langchain_core.callbacks.manager import CallbackManagerForLLMRun
from langchain_core.documents import Document
from langchain_core.language_models.llms import LLM
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langgraph.graph import END, StateGraph

import db
import memory_cache

try:
    from ddgs import DDGS
except ImportError:
    from duckduckgo_search import DDGS


# ===========================================================================
# 1. Logging
# ===========================================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("rag_chatbot.log", encoding="utf-8"),
    ],
)
logger = logging.getLogger("rag_backend")


# ===========================================================================
# 2. Config
# ===========================================================================
class Config:
    NGROK_URL: str = os.getenv("NGROK_URL", "")
    APP_API_KEY: str = os.getenv("APP_API_KEY", "")

    EMBEDDING_MODEL_NAME: str = "sentence-transformers/all-MiniLM-L6-v2"
    FAISS_BASE_DIR: str = "faiss_indexes"

    # Chunking
    CHUNK_SIZE: int = 600
    CHUNK_OVERLAP: int = 80
    SEMANTIC_MAX_SENTENCES: int = 8
    SEMANTIC_MIN_CHARS: int = 120
    SEMANTIC_SIMILARITY_THRESHOLD: float = 0.55

    # Retrieval
    TOP_K: int = 3
    # FAISS L2 distance threshold. Lower = more similar.
    # 0.85 requires a meaningfully close match before retrieval is "strong".
    SIMILARITY_SCORE_THRESHOLD: float = 0.85
    OCR_TEXT_LENGTH_THRESHOLD: int = 20

    # LLM
    MAX_NEW_TOKENS: int = 512
    TEMPERATURE: float = 0.2
    LLM_REQUEST_TIMEOUT: int = 120

    # Web search
    WEB_SEARCH_MAX_RESULTS: int = 3
    SCRAPE_TIMEOUT: int = 10
    SCRAPE_MAX_CHARS: int = 2000

    # Memory
    # v4: SHORT_TERM_WINDOW replaced by SHORT_TERM_CHAR_BUDGET.
    # Budget-aware pair selection packs as many recent Q&A pairs as fit
    # within SHORT_TERM_CHAR_BUDGET chars (matches build_prompt's 1200-char
    # cap for the short-term block). At ~450-500 chars/pair this yields
    # 2-3 pairs on typical conversations; more on short exchanges.
    SHORT_TERM_CHAR_BUDGET: int = 1200
    SUMMARY_CHAR_BUDGET: int = 600
    ENABLE_SUMMARY_MEMORY: bool = True
    SUMMARY_UPDATE_INTERVAL: int = 10

    # F8: Token budget (~1750 tokens safe for Qwen2.5-3B on T4)
    MAX_PROMPT_CHARS: int = 7000


# ===========================================================================
# 3. Groot System Prompt (F4)
# ===========================================================================
SYSTEM_PROMPT = """You are Groot — a friendly, intelligent, and helpful AI assistant.

Your personality:
- Warm and conversational, but always professional when the topic demands it.
- Occasionally witty — a well-placed joke or reference is welcome, but NEVER at the cost of accuracy.
- You may sometimes reference cricket, Bollywood, Indian culture, or internet memes — but only when it fits naturally. Do NOT force it.
- When someone asks "Who are you?", introduce yourself as Groot and briefly explain what you can do.

Your rules:
1. Accuracy first. Never sacrifice correctness for humor.
2. If you don't know something, say so honestly. Do not hallucinate.
3. Use the provided context (documents, web results, conversation history) to answer grounded in facts.
4. Keep answers concise unless the user explicitly asks for detail.
5. For technical questions, be precise and structured.
6. If humor doesn't fit (medical, legal, emotional topics), skip it entirely.

You have access to:
- Uploaded documents (PDFs, Word files, spreadsheets, images)
- Web search results (when documents don't have the answer)
- Full conversation history
"""


# ===========================================================================
# 4. Qwen LLM Wrapper
# ===========================================================================
class QwenNgrokLLM(LLM):
    """LangChain-compatible wrapper for Qwen2.5-3B-Instruct on Kaggle/ngrok."""

    ngrok_url: str
    api_key: str
    max_new_tokens: int = Config.MAX_NEW_TOKENS
    temperature: float = Config.TEMPERATURE
    timeout: int = Config.LLM_REQUEST_TIMEOUT

    @property
    def _llm_type(self) -> str:
        return "qwen-ngrok"

    def _call(
        self,
        prompt: str,
        stop: Optional[List[str]] = None,
        run_manager: Optional[CallbackManagerForLLMRun] = None,
        **kwargs: Any,
    ) -> str:
        logger.info("Calling Qwen LLM (prompt=%d chars)", len(prompt))
        try:
            response = requests.post(
                f"{self.ngrok_url}/generate",
                json={
                    "prompt": prompt,
                    "max_new_tokens": kwargs.get("max_new_tokens", self.max_new_tokens),
                    "temperature": kwargs.get("temperature", self.temperature),
                },
                headers={"x-api-key": self.api_key},
                timeout=self.timeout,
            )
            response.raise_for_status()
            text = response.json()["response"]
            logger.info("LLM responded (%d chars)", len(text))
            return text
        except requests.exceptions.RequestException as e:
            logger.error("LLM request failed: %s", e)
            raise RuntimeError(f"LLM unreachable: {e}") from e

    @property
    def _identifying_params(self) -> Mapping[str, Any]:
        return {"ngrok_url": self.ngrok_url, "max_new_tokens": self.max_new_tokens}


def create_llm() -> QwenNgrokLLM:
    if not Config.NGROK_URL or not Config.APP_API_KEY:
        logger.warning(
            "NGROK_URL and/or APP_API_KEY not set — chat will fail until configured in .env."
        )
    return QwenNgrokLLM(
        ngrok_url=Config.NGROK_URL,
        api_key=Config.APP_API_KEY,
        max_new_tokens=Config.MAX_NEW_TOKENS,
        temperature=Config.TEMPERATURE,
    )


def check_llm_health() -> bool:
    if not Config.NGROK_URL:
        return False
    try:
        r = requests.get(f"{Config.NGROK_URL}/health", timeout=10)
        return r.status_code == 200
    except requests.exceptions.RequestException as e:
        logger.warning("LLM health check failed: %s", e)
        return False


# ===========================================================================
# 5. Document Loading & Text Extraction
# ===========================================================================
def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    num_pages = len(reader.pages)
    text_parts = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(text_parts).strip()
    avg_chars = len(text) / max(num_pages, 1)
    if avg_chars < Config.OCR_TEXT_LENGTH_THRESHOLD:
        logger.info("PDF appears scanned (avg %.1f chars/page) — using OCR", avg_chars)
        return ocr_pdf(file_path)
    logger.info("Extracted %d chars from PDF text layer", len(text))
    return text


def ocr_pdf(file_path: str) -> str:
    try:
        images = convert_from_path(
            file_path,
            poppler_path=r"C:\poppler\poppler-26.02.0\Library\bin",
        )
    except Exception as e:
        logger.error("PDF→image conversion failed: %s", e)
        raise
    ocr_text = []
    for i, image in enumerate(images, start=1):
        logger.info("OCR page %d/%d", i, len(images))
        ocr_text.append(pytesseract.image_to_string(image))
    return "\n".join(ocr_text).strip()


def extract_text_from_image(file_path: str) -> str:
    logger.info("OCR on image: %s", file_path)
    return pytesseract.image_to_string(Image.open(file_path)).strip()


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    logger.info("Extracted %d chars from DOCX", len(text))
    return text


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    logger.info("Read %d chars from TXT", len(text))
    return text


def load_document(file_path: str, file_name: str) -> str:
    ext = os.path.splitext(file_name)[1].lower()
    logger.info("Loading '%s' (type: %s)", file_name, ext)
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        return extract_text_from_txt(file_path)
    elif ext in (".png", ".jpg", ".jpeg"):
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ===========================================================================
# 6. Semantic Chunking with Embeddings — used for ALL document types (F3)
# ===========================================================================
_SECTION_HEADER_RE = re.compile(
    r"^(?:"
    r"\d+(?:\.\d+)*[\s\.\)]+[A-Z]"   # "1. Title" or "2.1 Section"
    r"|[A-Z][A-Z\s]{4,}$"            # ALL CAPS line (≥5 chars)
    r"|.{3,60}:$"                     # "Some label:" at line end
    r")",
    re.MULTILINE,
)

_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+(?=[A-Z])")


def _split_sentences(text: str) -> List[str]:
    sentences = _SENTENCE_SPLIT_RE.split(text)
    result = []
    for s in sentences:
        for line in s.split("\n"):
            line = line.strip()
            if line:
                result.append(line)
    return result


def _is_section_header(line: str) -> bool:
    return bool(_SECTION_HEADER_RE.match(line.strip()))


def _cosine_sim(a: np.ndarray, b: np.ndarray) -> float:
    """Cosine similarity between two vectors, safe for zero vectors."""
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def _semantic_group_sentences(
    sentences: List[str],
    similarity_threshold: float = Config.SEMANTIC_SIMILARITY_THRESHOLD,
) -> List[List[str]]:
    """
    Group sentences into semantically coherent runs using embedding
    similarity between consecutive sentences. A new group starts when:
      1. Similarity drops below threshold (topic shift)
      2. Size/char caps are hit (prevent pathologically large chunks)
    """
    if not sentences:
        return []
    if len(sentences) == 1:
        return [sentences]

    embeddings = get_embeddings().embed_documents(sentences)
    embeddings = [np.array(e) for e in embeddings]

    groups: List[List[str]] = [[sentences[0]]]
    group_chars = len(sentences[0])

    for i in range(1, len(sentences)):
        sim = _cosine_sim(embeddings[i - 1], embeddings[i])
        sentence = sentences[i]
        exceeds_caps = (
            len(groups[-1]) >= Config.SEMANTIC_MAX_SENTENCES
            or group_chars + len(sentence) > Config.CHUNK_SIZE
        )
        if sim < similarity_threshold or exceeds_caps:
            groups.append([sentence])
            group_chars = len(sentence)
        else:
            groups[-1].append(sentence)
            group_chars += len(sentence)

    return groups


def semantic_chunk_text(text: str, source_name: str) -> List[Document]:
    """
    Split text into semantically coherent chunks using:
      1. Regex section-header detection (hard boundaries).
      2. Embedding-similarity grouping within each section (soft boundaries).

    Raises ValueError if 0 chunks are produced.
    """
    sentences = _split_sentences(text)
    if not sentences:
        raise ValueError(
            f"semantic_chunk_text: no sentences extracted from '{source_name}'. "
            "Check that the document has readable text content."
        )

    chunks: List[Document] = []
    current_section: str = "Introduction"
    section_sentences: List[str] = []

    def flush_section(section: str) -> None:
        nonlocal section_sentences
        if not section_sentences:
            return
        for group in _semantic_group_sentences(section_sentences):
            content = " ".join(group)
            if len(content) >= Config.SEMANTIC_MIN_CHARS:
                chunks.append(Document(
                    page_content=content,
                    metadata={"source": source_name, "section": section},
                ))
            elif chunks:
                chunks[-1] = Document(
                    page_content=chunks[-1].page_content + " " + content,
                    metadata=chunks[-1].metadata,
                )
            else:
                chunks.append(Document(
                    page_content=content,
                    metadata={"source": source_name, "section": section},
                ))
        section_sentences = []

    for sentence in sentences:
        if _is_section_header(sentence):
            flush_section(current_section)
            current_section = sentence.strip()
            continue
        section_sentences.append(sentence)

    flush_section(current_section)

    if not chunks:
        raise ValueError(
            f"semantic_chunk_text: produced 0 chunks from '{source_name}'. "
            "The document may be too short or contain only non-sentence content."
        )

    logger.info("Semantic chunking: '%s' → %d chunk(s)", source_name, len(chunks))
    return chunks


# ===========================================================================
# 7. Tabular Ingestion — Excel & CSV
# ===========================================================================
def extract_rows_from_excel(file_path: str) -> Tuple[List[Dict[str, Any]], Dict[str, pd.DataFrame]]:
    logger.info("Opening Excel workbook: %s", file_path)
    workbook = load_workbook(file_path, read_only=True, data_only=True)
    rows_out: List[Dict[str, Any]] = []
    dataframes: Dict[str, pd.DataFrame] = {}

    for sheet in workbook.worksheets:
        sheet_name = sheet.title
        rows_iter = sheet.iter_rows(values_only=True)
        try:
            header = next(rows_iter)
        except StopIteration:
            logger.warning("Sheet '%s' is empty — skipping", sheet_name)
            continue

        headers = [
            str(h).strip() if h is not None else f"col_{i}"
            for i, h in enumerate(header)
        ]
        sheet_rows = []
        for row_idx, row in enumerate(rows_iter, start=2):
            if row is None or all(cell is None for cell in row):
                continue
            row_dict = {}
            pairs = []
            for col, val in zip(headers, row):
                if val is not None and str(val).strip():
                    pairs.append(f"{col} is {val}")
                    row_dict[col] = val
            if not pairs:
                continue
            rows_out.append({
                "sheet": sheet_name,
                "row_index": row_idx,
                "text": f"In sheet '{sheet_name}', row {row_idx}: " + "; ".join(pairs) + ".",
                "row_data": row_dict,
            })
            sheet_rows.append(row_dict)

        if sheet_rows:
            dataframes[sheet_name] = pd.DataFrame(sheet_rows)

    workbook.close()
    logger.info("Excel: %d row(s) extracted across %d sheet(s)", len(rows_out), len(dataframes))
    return rows_out, dataframes


def extract_rows_from_csv(
    file_path: str, read_chunksize: int = 5000
) -> Tuple[List[Dict[str, Any]], Dict[str, pd.DataFrame]]:
    rows_out: List[Dict[str, Any]] = []
    full_df_parts: List[pd.DataFrame] = []

    for batch_idx, df_chunk in enumerate(
        pd.read_csv(file_path, chunksize=read_chunksize), start=1
    ):
        full_df_parts.append(df_chunk)
        columns = [str(c).strip() for c in df_chunk.columns]
        for idx, row in df_chunk.iterrows():
            pairs = [
                f"{col} is {row[col]}"
                for col in columns
                if pd.notna(row[col]) and str(row[col]).strip()
            ]
            if not pairs:
                continue
            rows_out.append({
                "sheet": "csv",
                "row_index": int(idx),
                "text": f"Row {idx}: " + "; ".join(pairs) + ".",
            })
        logger.info("CSV batch %d: %d row(s) total", batch_idx, len(rows_out))

    dataframes = {"csv": pd.concat(full_df_parts, ignore_index=True)} if full_df_parts else {}
    logger.info("CSV: %d row(s) extracted", len(rows_out))
    return rows_out, dataframes


def chunk_tabular_rows_semantic(
    rows: List[Dict[str, Any]],
    source_name: str,
) -> List[Document]:
    if not rows:
        return []

    row_texts = [r["text"] for r in rows]
    groups = _semantic_group_sentences(row_texts, Config.SEMANTIC_SIMILARITY_THRESHOLD)

    documents: List[Document] = []
    row_cursor = 0
    for group in groups:
        n = len(group)
        batch = rows[row_cursor : row_cursor + n]
        row_cursor += n
        documents.append(Document(
            page_content="\n".join(r["text"] for r in batch),
            metadata={
                "source": source_name,
                "sheets": ", ".join(sorted({r["sheet"] for r in batch})),
                "row_start": batch[0]["row_index"],
                "row_end": batch[-1]["row_index"],
            },
        ))

    logger.info(
        "Tabular semantic chunking: %d row(s) → %d chunk(s) for '%s'",
        len(rows), len(documents), source_name,
    )
    return documents


# ---------------------------------------------------------------------------
# DataFrame store
# ---------------------------------------------------------------------------
_dataframe_store: Dict[str, Dict[str, pd.DataFrame]] = {}


def store_dataframes(conversation_id: str, dataframes: Dict[str, pd.DataFrame]) -> None:
    if conversation_id not in _dataframe_store:
        _dataframe_store[conversation_id] = {}
    _dataframe_store[conversation_id].update(dataframes)
    logger.info(
        "DataFrame store: conversation %s now has tables: %s",
        conversation_id,
        list(_dataframe_store[conversation_id].keys()),
    )


def process_uploaded_file(
    file_path: str,
    file_name: str,
    vectorstore: Optional[FAISS],
    conversation_id: str,
) -> Tuple[FAISS, int]:
    ext = os.path.splitext(file_name)[1].lower()

    if ext in (".xlsx", ".xls"):
        rows, dataframes = extract_rows_from_excel(file_path)
        if not rows:
            raise ValueError(f"No data rows extracted from '{file_name}'.")
        store_dataframes(conversation_id, dataframes)
        documents = chunk_tabular_rows_semantic(rows, file_name)

    elif ext == ".csv":
        rows, dataframes = extract_rows_from_csv(file_path)
        if not rows:
            raise ValueError(f"No data rows extracted from '{file_name}'.")
        store_dataframes(conversation_id, dataframes)
        documents = chunk_tabular_rows_semantic(rows, file_name)

    else:
        text = load_document(file_path, file_name)
        if not text.strip():
            raise ValueError(f"No text extracted from '{file_name}'.")
        documents = semantic_chunk_text(text, file_name)

    if not documents:
        raise ValueError(
            f"Chunking produced 0 documents for '{file_name}'. "
            "The file may be empty or contain unsupported content."
        )

    vectorstore = add_documents_to_vectorstore(documents, vectorstore)
    return vectorstore, len(documents)


# ===========================================================================
# 8. Tabular Analytics — SQL & Math Execution
# ===========================================================================
_SQL_EXTRACT_RE = re.compile(
    r"(SELECT\s+.+?)(?=\s*[;\"']|$)",
    re.IGNORECASE | re.DOTALL,
)

_MATH_AGG_RE = re.compile(
    r"\b(sum|total|average|avg|mean|count|how\s+many|max|maximum|highest|largest"
    r"|min|minimum|lowest|smallest|median|std|variance)\b",
    re.IGNORECASE,
)


def _safe_table_name(name: str) -> str:
    return re.sub(r"[^\w]", "_", name).lower().strip("_") or "data"


def get_dataframe_schema(conversation_id: str) -> str:
    dfs = _dataframe_store.get(conversation_id, {})
    if not dfs:
        return ""
    parts = []
    for name, df in dfs.items():
        safe = _safe_table_name(name)
        col_info = ", ".join(
            f"{c} ({df[c].dtype})" for c in df.columns
        )
        parts.append(f"Table '{safe}' ({len(df)} rows): {col_info}")
    return "\n".join(parts)


def _extract_sql_from_question(question: str) -> Optional[str]:
    match = _SQL_EXTRACT_RE.search(question)
    return match.group(1).strip() if match else None


def _find_best_column(question: str, df: pd.DataFrame) -> Optional[str]:
    q_lower = question.lower()
    cols = df.columns.tolist()

    for col in cols:
        if col.lower() in q_lower:
            return col

    q_words = set(re.findall(r"\w+", q_lower))
    best_col, best_overlap = None, 0
    for col in cols:
        col_words = set(re.findall(r"\w+", col.lower()))
        overlap = len(col_words & q_words)
        if overlap > best_overlap:
            best_col, best_overlap = col, overlap

    return best_col if best_overlap > 0 else None


def execute_sql_on_dataframes(
    sql_query: str, conversation_id: str
) -> Tuple[str, bool]:
    dfs = _dataframe_store.get(conversation_id, {})
    if not dfs:
        return "No tabular data is loaded for this conversation.", False

    conn = sqlite3.connect(":memory:")
    try:
        for name, df in dfs.items():
            safe = _safe_table_name(name)
            df.to_sql(safe, conn, if_exists="replace", index=False)
            logger.info("SQL: registered table '%s' (%d rows)", safe, len(df))

        result_df = pd.read_sql_query(sql_query, conn)

        if result_df.empty:
            return "The query returned no results.", True

        if len(result_df) > 50:
            result_text = (
                result_df.head(50).to_string(index=False)
                + f"\n\n[Output truncated — showing 50 of {len(result_df)} rows]"
            )
        else:
            result_text = result_df.to_string(index=False)

        logger.info("SQL execution succeeded: %d row(s) returned", len(result_df))
        return result_text, True

    except Exception as e:
        logger.error("SQL execution error: %s", e)
        return f"SQL execution error: {e}", False
    finally:
        conn.close()


def execute_math_on_dataframes(
    question: str, conversation_id: str
) -> Optional[str]:
    dfs = _dataframe_store.get(conversation_id, {})
    if not dfs:
        return None

    target_df: Optional[pd.DataFrame] = None
    target_table: str = ""
    for name, df in dfs.items():
        if name.lower() in question.lower() or _safe_table_name(name) in question.lower():
            target_df, target_table = df, name
            break
    if target_df is None:
        target_table, target_df = next(iter(dfs.items()))

    q = question.lower()
    col = _find_best_column(question, target_df)

    results: List[str] = []

    try:
        if re.search(r"\b(sum|total)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                results.append(f"Sum of '{col}': {target_df[col].sum():,.4f}")
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column sums:\n" + num_cols.sum().to_string())

        if re.search(r"\b(avg|average|mean)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                results.append(f"Average of '{col}': {target_df[col].mean():,.4f}")
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column averages:\n" + num_cols.mean().to_string())

        if re.search(r"\b(count|how\s+many)\b", q):
            results.append(f"Total row count in '{target_table}': {len(target_df):,}")
            if col:
                non_null = target_df[col].notna().sum()
                results.append(f"Non-null count in '{col}': {non_null:,}")

        if re.search(r"\b(max|maximum|highest|largest)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                idx = target_df[col].idxmax()
                results.append(
                    f"Maximum of '{col}': {target_df[col].max():,.4f} (row {idx})"
                )
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column maximums:\n" + num_cols.max().to_string())

        if re.search(r"\b(min|minimum|lowest|smallest)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                idx = target_df[col].idxmin()
                results.append(
                    f"Minimum of '{col}': {target_df[col].min():,.4f} (row {idx})"
                )
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column minimums:\n" + num_cols.min().to_string())

        if re.search(r"\b(median)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                results.append(f"Median of '{col}': {target_df[col].median():,.4f}")
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column medians:\n" + num_cols.median().to_string())

        if re.search(r"\b(std|standard\s+deviation|variance)\b", q):
            if col and pd.api.types.is_numeric_dtype(target_df[col]):
                results.append(
                    f"Std dev of '{col}': {target_df[col].std():,.4f} | "
                    f"Variance: {target_df[col].var():,.4f}"
                )
            else:
                num_cols = target_df.select_dtypes(include="number")
                results.append("Column std devs:\n" + num_cols.std().to_string())

    except Exception as e:
        logger.error("Math execution error: %s", e)
        return f"Math computation error: {e}"

    if not results:
        return None

    header = f"[Table: '{target_table}'" + (f", Column: '{col}'" if col else "") + "]"
    return header + "\n" + "\n".join(results)


def run_tabular_analytics(
    question: str, conversation_id: str
) -> Tuple[str, bool]:
    # 1. Explicit SQL
    sql = _extract_sql_from_question(question)
    if sql:
        logger.info("Tabular analytics: explicit SQL detected")
        result, ok = execute_sql_on_dataframes(sql, conversation_id)
        return f"SQL Query:\n{sql}\n\nResult:\n{result}", ok

    # 2. NL math/aggregation
    math_result = execute_math_on_dataframes(question, conversation_id)
    if math_result:
        logger.info("Tabular analytics: math operation mapped")
        return f"Computed Result:\n{math_result}", True

    # 3. Fallback: schema context for LLM
    schema = get_dataframe_schema(conversation_id)
    fallback = (
        "No specific computation could be mapped to this question.\n\n"
        f"Available table schema:\n{schema}"
        if schema
        else "No tabular data loaded for this conversation."
    )
    logger.info("Tabular analytics: no mapping found, returning schema")
    return fallback, False


# ===========================================================================
# 9. FAISS Vector Store
# ===========================================================================
_embeddings_instance: Optional[HuggingFaceEmbeddings] = None


def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings_instance
    if _embeddings_instance is None:
        logger.info("Loading embedding model: %s", Config.EMBEDDING_MODEL_NAME)
        _embeddings_instance = HuggingFaceEmbeddings(
            model_name=Config.EMBEDDING_MODEL_NAME
        )
    return _embeddings_instance


def get_vectorstore_path(conversation_id: str) -> str:
    return os.path.join(Config.FAISS_BASE_DIR, str(conversation_id))


def add_documents_to_vectorstore(
    documents: List[Document],
    vectorstore: Optional[FAISS],
) -> FAISS:
    embeddings = get_embeddings()
    if vectorstore is None:
        logger.info("Creating new FAISS index (%d chunks)", len(documents))
        return FAISS.from_documents(documents, embeddings)
    logger.info("Adding %d chunks to FAISS index", len(documents))
    vectorstore.add_documents(documents)
    return vectorstore


def save_vectorstore(vectorstore: FAISS, path: str) -> None:
    parent = os.path.dirname(path)
    if parent:
        os.makedirs(parent, exist_ok=True)
    vectorstore.save_local(path)
    logger.info("Vector store saved: '%s'", path)


def load_vectorstore(path: str) -> Optional[FAISS]:
    if not os.path.exists(path):
        return None
    logger.info("Loading vector store: '%s'", path)
    return FAISS.load_local(
        path, get_embeddings(), allow_dangerous_deserialization=True
    )


def get_vectorstore_status(
    vectorstore: Optional[FAISS],
    faiss_path: str,
    conversation_id: str,
) -> Dict[str, Any]:
    chunk_count = vectorstore.index.ntotal if vectorstore is not None else 0
    try:
        document_count = len(db.list_documents(conversation_id))
    except Exception as e:
        logger.error("Could not fetch document count: %s", e)
        document_count = 0
    return {
        "loaded": vectorstore is not None,
        "embedding_model": Config.EMBEDDING_MODEL_NAME,
        "document_count": document_count,
        "chunk_count": chunk_count,
        "faiss_path": faiss_path,
        "analytics_tables": list(_dataframe_store.get(conversation_id, {}).keys()),
    }


# ===========================================================================
# 10. Web Search + Scraping
# ===========================================================================
def web_search(
    query: str, max_results: int = Config.WEB_SEARCH_MAX_RESULTS
) -> List[Dict[str, str]]:
    logger.info("Web search: %s", query)
    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))
        logger.info("Web search: %d result(s)", len(results))
        return results
    except Exception as e:
        logger.error("Web search failed: %s", e)
        return []


def scrape_url(url: str, timeout: int = Config.SCRAPE_TIMEOUT) -> str:
    logger.info("Scraping: %s", url)
    try:
        response = requests.get(
            url, timeout=timeout, headers={"User-Agent": "Mozilla/5.0"}
        )
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return " ".join(soup.stripped_strings)[: Config.SCRAPE_MAX_CHARS]
    except requests.exceptions.RequestException as e:
        logger.warning("Scrape failed for %s: %s", url, e)
        return ""


def gather_web_context(query: str) -> str:
    results = web_search(query)
    if not results:
        return ""
    parts = []
    for result in results:
        url = result.get("href", "")
        content = scrape_url(url) if url else result.get("body", "")
        if content:
            parts.append(f"Source: {url}\n{content}")
    return "\n\n".join(parts)


# ===========================================================================
# 11. Retrieval Quality & Routing
# ===========================================================================

# v4 fix: removed stale 'follow_up_query' (intent no longer exists since v3
# split it into follow_up_doc / follow_up_chat). Added follow_up_chat,
# follow_up_doc, and knowledge_query — all three should never hit web search.
_NO_WEB_SEARCH_INTENTS = frozenset({
    "general_chat",      # greetings / identity — no external lookup needed
    "follow_up_chat",    # memory-only follow-up — no docs, no web
    "follow_up_doc",     # follow-up with docs loaded — trust retrieval, not web
    "tabular_analytics", # computed locally from stored DataFrames
    "knowledge_query",   # LLM answers from training — web adds latency, no benefit
})


def assess_retrieval_quality(
    vectorstore: Optional[FAISS],
    query: str,
    k: int = Config.TOP_K,
) -> Tuple[List[Document], bool]:
    if vectorstore is None:
        return [], False
    try:
        results = vectorstore.similarity_search_with_score(query, k=k)
    except Exception as e:
        logger.error("Retrieval failed: %s", e)
        return [], False
    if not results:
        return [], False
    docs = [doc for doc, _ in results]
    best_score = min(score for _, score in results)
    is_strong = best_score <= Config.SIMILARITY_SCORE_THRESHOLD
    logger.info("Retrieval: best_score=%.4f strong=%s", best_score, is_strong)
    return docs, is_strong


def should_use_web_search(
    vectorstore: Optional[FAISS],
    retrieval_strong: bool,
    intent: str,
    doc_context: str,
) -> bool:
    """
    Intent-aware web search gate.

    Rules (in priority order):
      1. Never search for no-web intents (see _NO_WEB_SEARCH_INTENTS).
      2. Never search for document_query when we already have doc context.
      3. Always search for explicit web_search_query intent.
      4. Search for technical_question when no vectorstore OR retrieval weak.
      5. Otherwise don't search.
    """
    if intent in _NO_WEB_SEARCH_INTENTS:
        logger.info("Web search blocked: intent=%s is in no-web list", intent)
        return False

    if intent == "document_query" and doc_context.strip():
        logger.info("Web search blocked: document_query with existing doc context")
        return False

    if intent == "web_search_query":
        logger.info("Web search triggered: explicit web_search_query intent")
        return True

    if vectorstore is None:
        logger.info("Web search triggered: no vectorstore")
        return True

    if not retrieval_strong:
        logger.info(
            "Web search triggered: retrieval not strong (score > %.2f)",
            Config.SIMILARITY_SCORE_THRESHOLD,
        )
        return True

    logger.info("Web search skipped: strong retrieval exists")
    return False


# ===========================================================================
# 12. Follow-up Detection (F6)
# ===========================================================================
_FOLLOWUP_TRIGGERS = re.compile(
    r"^("
    r"(what|how|why|when|where|who|can|could|should|would|does|is|are|do|did)\s+(it|this|that|they|he|she|we)\b"
    r"|tell\s+me\s+more"
    r"|explain\s+(that|this|it|more)"
    r"|what\s+about"
    r"|and\s+(what|how|why|can|does)"
    r"|more\s+(details?|info|examples?)"
    r"|elaborate|go\s+on|continue"
    r"|so\s+(it|this|that|they)"
    r")",
    re.IGNORECASE,
)


def is_follow_up(question: str) -> bool:
    return bool(_FOLLOWUP_TRIGGERS.match(question.strip()))


def expand_query_with_context(question: str, recent_messages: list) -> str:
    if not recent_messages or len(recent_messages) < 2:
        return question
    for m in reversed(recent_messages[:-1]):
        role = m.get("role") if isinstance(m, dict) else m.role
        if role == "user":
            prior = m.get("content") if isinstance(m, dict) else m.content
            expanded = f"{prior[:200]} {question}"
            logger.info("Follow-up expanded query: %s", expanded[:120])
            return expanded
    return question


# ===========================================================================
# 13. Conversation Memory (F1)
#     v4: budget-aware pair selection + incremental summary
#
#     Design:
#       - _filter_relevant_turns REMOVED — was broken (referenced undefined
#         Config constants RELEVANT_TURNS_TOP_K and ALWAYS_INCLUDE_LAST_PAIRS,
#         causing AttributeError at import time) and was over-filtering context
#         the LLM needed.
#       - SHORT_TERM_WINDOW removed from Config — replaced by SHORT_TERM_CHAR_BUDGET.
#       - New _select_pairs_within_budget:
#           * Builds (user, assistant) pairs from full message list.
#           * Iterates newest → oldest, packing into SHORT_TERM_CHAR_BUDGET chars.
#           * Returns pairs in original chronological order for the LLM.
#           * Dynamic count: 2-3 pairs on typical exchanges; more on short ones.
#       - maybe_update_long_term_summary: INCREMENTAL — only the last
#         SUMMARY_UPDATE_INTERVAL new turns are summarized and merged with
#         the prior summary. Never re-summarizes the full history.
#       - build_memory_context: fetches up to 20 messages (10 pairs), budget
#         selector trims to what fits.
# ===========================================================================
@dataclass
class MemoryMessage:
    """Unified message format from either Redis (dict) or Postgres (ORM)."""
    role: str
    content: str


def _build_pairs(
    messages: List[MemoryMessage],
) -> List[Tuple[MemoryMessage, Optional[MemoryMessage]]]:
    """
    Convert a flat message list into (user, assistant) pairs.
    A user message with no following assistant reply gets (user, None).
    Unpaired leading assistant messages are skipped.
    """
    pairs: List[Tuple[MemoryMessage, Optional[MemoryMessage]]] = []
    i = 0
    while i < len(messages):
        if messages[i].role == "user":
            if i + 1 < len(messages) and messages[i + 1].role == "assistant":
                pairs.append((messages[i], messages[i + 1]))
                i += 2
            else:
                pairs.append((messages[i], None))
                i += 1
        else:
            # Skip unpaired assistant messages
            i += 1
    return pairs


def _pair_to_text(
    pair: Tuple[MemoryMessage, Optional[MemoryMessage]],
) -> str:
    """Render a (user, assistant) pair as Q&A text for the prompt."""
    user_msg, asst_msg = pair
    lines = [f"Q: {user_msg.content}"]
    if asst_msg is not None:
        lines.append(f"A: {asst_msg.content}")
    return "\n".join(lines)


def _select_pairs_within_budget(
    messages: List[MemoryMessage],
    char_budget: int = Config.SHORT_TERM_CHAR_BUDGET,
) -> str:
    """
    Select the most recent Q&A pairs that fit within char_budget chars.

    Strategy:
      1. Build (user, assistant) pairs from the full message list.
      2. Iterate newest → oldest, accumulating pairs until budget is exhausted.
      3. Return selected pairs in original chronological order (oldest first)
         so the LLM reads them in the correct sequence.

    Edge case: if a single pair exceeds the budget (very long assistant reply),
    it is included truncated — the LLM must always have at least the most
    recent exchange for coherent follow-up handling.

    Replaces the old fixed SHORT_TERM_WINDOW and the broken
    _filter_relevant_turns function.
    """
    pairs = _build_pairs(messages)
    if not pairs:
        return ""

    selected: List[Tuple[MemoryMessage, Optional[MemoryMessage]]] = []
    chars_used = 0

    # Newest → oldest to respect recency priority
    for pair in reversed(pairs):
        text = _pair_to_text(pair)
        pair_chars = len(text) + 2  # +2 for "\n\n" separator
        if chars_used + pair_chars > char_budget:
            break
        selected.append(pair)
        chars_used += pair_chars

    if not selected:
        # Single pair exceeds budget — include it truncated so the LLM
        # always has at least the most recent exchange.
        last_pair = pairs[-1]
        text = _pair_to_text(last_pair)
        truncated = text[: char_budget - 20] + "\n[... truncated ...]"
        logger.info(
            "Memory: single pair exceeded budget (%d chars), truncated to %d chars",
            len(text), char_budget,
        )
        return truncated

    # Reverse back to chronological order for the LLM
    selected.reverse()

    result = "\n\n".join(_pair_to_text(p) for p in selected)
    logger.info(
        "Memory: selected %d/%d pairs within %d-char budget (%d chars used)",
        len(selected), len(pairs), char_budget, chars_used,
    )
    return result


def _format_messages_as_qa(messages: List[MemoryMessage]) -> str:
    """
    Format a flat message list as Q&A text.

    Used by maybe_update_long_term_summary to format only the NEW turns
    since the last summary update — not the full history.
    """
    lines: List[str] = []
    i = 0
    while i < len(messages):
        m = messages[i]
        content = m.content[:600]
        if m.role == "user":
            lines.append(f"Q: {content}")
            if i + 1 < len(messages) and messages[i + 1].role == "assistant":
                a_content = messages[i + 1].content[:600]
                lines.append(f"A: {a_content}")
                i += 2
                continue
        else:
            lines.append(f"A: {content}")
        i += 1
    return "\n\n".join(lines)


def build_memory_context(
    conversation_id: str, question: str = ""
) -> Tuple[str, str, list]:
    """
    Returns: (short_term_text, long_term_summary, raw_messages)

    v4 changes:
    - Fetches up to 20 messages (10 pairs) from Redis/Postgres.
    - _select_pairs_within_budget trims to SHORT_TERM_CHAR_BUDGET chars,
      always keeping the most recent pairs.
    - No relevance filtering (_filter_relevant_turns removed — was broken).
    - raw_messages still returned for follow-up detection downstream.

    Tries Redis first; falls back to Postgres on miss.
    """
    cached = memory_cache.get_recent_messages(conversation_id)

    if cached is not None:
        messages = [
            MemoryMessage(role=m["role"], content=m["content"]) for m in cached
        ]
    else:
        # Fetch 20 messages = up to 10 pairs; budget selector will trim to fit.
        db_messages = db.get_messages(conversation_id, limit=20)
        messages = [
            MemoryMessage(role=m.role, content=m.content) for m in db_messages
        ]

    short_term = _select_pairs_within_budget(messages, Config.SHORT_TERM_CHAR_BUDGET)

    convo = db.get_conversation(conversation_id)
    long_term_summary = (convo.summary or "") if convo else ""

    return short_term, long_term_summary, messages


def maybe_update_long_term_summary(
    llm: QwenNgrokLLM, conversation_id: str
) -> None:
    """
    v4: incremental summary — only the NEW turns since the last summary
    update are summarized and merged into the prior summary.

    Fires every SUMMARY_UPDATE_INTERVAL messages (same trigger as before).
    Only the last SUMMARY_UPDATE_INTERVAL messages are sent to the LLM,
    keeping each summary generation cheap regardless of conversation length.

    The LLM prompt merges prior summary + new turns → updated summary
    without re-reading old turns verbatim.
    """
    if not Config.ENABLE_SUMMARY_MEMORY:
        return

    total = db.count_messages(conversation_id)
    if total == 0 or total % Config.SUMMARY_UPDATE_INTERVAL != 0:
        return

    convo = db.get_conversation(conversation_id)
    if convo is None:
        return

    prior_summary = convo.summary or ""

    # Fetch ONLY the last SUMMARY_UPDATE_INTERVAL messages (the new turns
    # since the previous summary update). This is the incremental part —
    # we never re-summarize the full history.
    new_messages = db.get_messages(
        conversation_id, limit=Config.SUMMARY_UPDATE_INTERVAL
    )
    new_turns_text = _format_messages_as_qa(
        [MemoryMessage(role=m.role, content=m.content) for m in new_messages]
    )

    if not new_turns_text.strip():
        logger.info(
            "Summary update skipped: no new turns to summarize for %s",
            conversation_id,
        )
        return

    # Incremental merge: prior summary + new turns → updated summary
    summary_prompt = (
        "You are maintaining a running summary of a conversation between a user "
        "and Groot (an AI assistant).\n\n"
        "You will receive:\n"
        "  1. The PRIOR SUMMARY — what was discussed before.\n"
        "  2. NEW TURNS — the most recent exchanges to incorporate.\n\n"
        "Your task: merge them into a single updated summary under 150 words that "
        "explicitly preserves:\n"
        "  - Key topics and questions discussed\n"
        "  - Specific facts, numbers, or data points mentioned\n"
        "  - Uploaded files and what was found or computed from them\n"
        "  - User's stated goals, preferences, or constraints\n"
        "  - Any open or unresolved questions\n\n"
        "Write ONLY the updated summary. No preamble, no labels, no explanation.\n\n"
        f"PRIOR SUMMARY:\n{prior_summary or '(none yet)'}\n\n"
        f"NEW TURNS:\n{new_turns_text}\n\n"
        "UPDATED SUMMARY:"
    )

    try:
        updated_summary = llm.invoke(summary_prompt)
        db.update_conversation_summary(conversation_id, updated_summary.strip())
        logger.info(
            "Incremental summary updated for %s (total messages: %d, new turns: %d)",
            conversation_id, total, len(new_messages),
        )
    except Exception as e:
        logger.error("Summary update failed for %s: %s", conversation_id, e)


# ===========================================================================
# 14. Title Generation (F5)
# ===========================================================================
def generate_conversation_title(
    llm: QwenNgrokLLM,
    user_question: str,
    assistant_answer: str,
) -> str:
    prompt = (
        "Generate a short conversation title (3–5 words, no punctuation) "
        "that captures the main topic of this exchange.\n\n"
        f"User: {user_question[:300]}\n"
        f"Assistant: {assistant_answer[:300]}\n\n"
        "Title (3-5 words only):"
    )
    try:
        title = llm.invoke(prompt, max_new_tokens=20, temperature=0.3).strip()
        title = title.split("\n")[0].strip('"\'').strip()
        if title and 3 <= len(title) <= 80:
            return title
    except Exception as e:
        logger.warning("Title generation failed: %s", e)

    words = user_question.strip().split()
    title = " ".join(words[:6])
    if len(words) > 6:
        title += "..."
    return title[:80]


def generate_conversation_title_simple(first_message: str) -> str:
    """Instant fallback title — no LLM call."""
    words = first_message.strip().split()
    if not words:
        return "New Chat"
    title = " ".join(words[:8])
    if len(words) > 8:
        title += "..."
    return title[:100]


# ===========================================================================
# 15. Prompt Builder (F7)
# ===========================================================================
def _trim_to_budget(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n[... truncated ...]"


def build_prompt(
    short_term_history: str,
    long_term_summary: str,
    doc_context: str,
    web_context: str,
    question: str,
) -> str:
    summary_block = (
        "### Conversation Summary (earlier context)\n"
        + _trim_to_budget(long_term_summary, Config.SUMMARY_CHAR_BUDGET)
        if long_term_summary
        else ""
    )
    memory_block = (
        "### Recent Conversation\n"
        + _trim_to_budget(short_term_history, Config.SHORT_TERM_CHAR_BUDGET)
        if short_term_history
        else ""
    )
    doc_block = (
        "### Document Context\n" + _trim_to_budget(doc_context, 2000)
        if doc_context
        else ""
    )
    web_block = (
        "### Web Search Results\n" + _trim_to_budget(web_context, 1200)
        if web_context
        else ""
    )

    context_parts = [b for b in [summary_block, memory_block, doc_block, web_block] if b]
    context_section = (
        "\n\n".join(context_parts) if context_parts else "No additional context available."
    )

    prompt = (
        f"{SYSTEM_PROMPT}\n\n"
        f"{context_section}\n\n"
        f"### Current Question\n"
        f"User: {question[:300]}\n"
        f"Groot:"
    )
    logger.info("Prompt built: %d chars", len(prompt))
    logger.info("=" * 80)
    logger.info(prompt)
    logger.info("=" * 80)
    return prompt


# ===========================================================================
# 16. LangGraph Pipeline (F2)
# ===========================================================================

# ---------------------------------------------------------------------------
# Graph State
# ---------------------------------------------------------------------------
class GraphState(TypedDict):
    question: str
    conversation_id: str
    vectorstore: Optional[FAISS]
    llm: Any

    intent: str
    retrieval_query: str
    is_follow_up: bool

    short_term: str
    long_term_summary: str
    raw_messages: list

    doc_context: str
    retrieval_strong: bool
    web_context: str

    answer: str
    sources: List[str]


# ---------------------------------------------------------------------------
# Intent Classification — rule-based (no extra LLM call)
# ---------------------------------------------------------------------------

# Tight — greetings and identity questions ONLY.
# "ok/sure/yes/no/got it" intentionally excluded: they are acknowledgements
# that may have a follow-up embedded and routing them to general_chat_node
# (zero retrieval context) produced bad answers.
_GENERAL_CHAT_RE = re.compile(
    r"^(hi|hello|hey|howdy)\b"
    r"|\b(what('s| is) your name|who are you|introduce yourself)\b"
    r"|^(thanks?|thank you|bye|goodbye)\b"
    r"|^good (morning|afternoon|evening)\b"
    r"|\bhow are you\b"
    r"|\b(my name is|i am|i'm|call me)\s+\w+"
    r"|what can you do|help me",
    re.IGNORECASE,
)

_WEB_SEARCH_SIGNALS = re.compile(
    r"\b(latest|current|recent|today|news|2024|2025|live|real.?time|stock price"
    r"|weather|score|winner|just released|new version|update)\b",
    re.IGNORECASE,
)

_DOCUMENT_SIGNALS = re.compile(
    r"\b(document|file|uploaded|pdf|report|according to|in the|the document"
    r"|spreadsheet|table|row|column|sheet|attachment)\b",
    re.IGNORECASE,
)

_ANALYTICS_RE = re.compile(
    r"\b(select\s+|sum|total|average|avg|mean|count|max|maximum|minimum|min"
    r"|how\s+many|group\s+by|where\s+|filter|aggregate|calculate|compute"
    r"|sql|query|median|std|variance|highest|lowest|largest|smallest)\b",
    re.IGNORECASE,
)


def _classify_intent(
    question: str,
    vectorstore: Optional[FAISS],
    follow_up: bool,
    conversation_id: str,
) -> str:
    """
    Strict priority cascade — first match wins.

      1. general_chat      — greetings / identity (tight regex, no retrieval)
      2. tabular_analytics — SQL or NL math on loaded DataFrames
      3. follow_up_doc     — follow-up + docs loaded → memory + retrieval
      4. follow_up_chat    — follow-up + no docs    → memory only
      5. web_search_query  — explicit freshness signal
      6. document_query    — docs loaded OR explicit doc mention
      7. knowledge_query   — no docs, no web signals → pure LLM knowledge
    """
    q = question.strip()

    # 1. Pure chat
    if _GENERAL_CHAT_RE.search(q):
        return "general_chat"

    # 2. Analytics on loaded DataFrames
    if (
        conversation_id in _dataframe_store
        and _dataframe_store[conversation_id]
        and _ANALYTICS_RE.search(q)
    ):
        return "tabular_analytics"

    # 3 & 4. Follow-up — split by whether docs are available
    if follow_up:
        if vectorstore is not None:
            return "follow_up_doc"
        return "follow_up_chat"

    # 5. Explicit web signal
    if _WEB_SEARCH_SIGNALS.search(q):
        return "web_search_query"

    # 6. Document query
    if vectorstore is not None or _DOCUMENT_SIGNALS.search(q):
        return "document_query"

    # 7. General knowledge
    return "knowledge_query"


# ---------------------------------------------------------------------------
# Nodes
# ---------------------------------------------------------------------------
def _intent_classifier_node(state: GraphState) -> GraphState:
    question = state["question"]
    follow_up = is_follow_up(question)
    state["is_follow_up"] = follow_up
    state["intent"] = _classify_intent(
        question,
        state.get("vectorstore"),
        follow_up,
        state["conversation_id"],
    )
    state["retrieval_query"] = question
    state.setdefault("short_term", "")
    state.setdefault("long_term_summary", "")
    state.setdefault("raw_messages", [])
    state.setdefault("doc_context", "")
    state.setdefault("retrieval_strong", False)
    state.setdefault("web_context", "")
    state.setdefault("sources", [])
    logger.info("Intent: %s | follow_up: %s", state["intent"], follow_up)
    return state


def _memory_node(state: GraphState) -> GraphState:
    try:
        short_term, long_term_summary, raw_messages = build_memory_context(
            state["conversation_id"], state["question"]
        )
        state["short_term"] = short_term
        state["long_term_summary"] = long_term_summary
        state["raw_messages"] = raw_messages
    except Exception as e:
        logger.error("Memory node failed: %s", e)
        state["short_term"] = ""
        state["long_term_summary"] = ""
        state["raw_messages"] = []

    if state["is_follow_up"] and state["raw_messages"]:
        state["retrieval_query"] = expand_query_with_context(
            state["question"], state["raw_messages"]
        )
    return state


def _memory_router(state: GraphState) -> str:
    intent = state["intent"]
    if intent == "general_chat":
        return "general_chat_node"
    if intent == "web_search_query":
        return "web_search_node"
    if intent == "tabular_analytics":
        return "tabular_analytics_node"
    if intent == "follow_up_chat":
        # v4 fix: no docs loaded — memory context alone is sufficient.
        # Skip retrieval entirely; go straight to response via general_chat_node.
        # (general_chat_node is a no-op pass-through that leads to response_node.)
        return "general_chat_node"
    # follow_up_doc, document_query, knowledge_query all go through retrieval.
    # For knowledge_query with no vectorstore, retrieval returns empty context
    # and should_use_web_search returns False (knowledge_query is in
    # _NO_WEB_SEARCH_INTENTS), so the LLM answers from its own knowledge.
    return "document_retrieval_node"


def _document_retrieval_node(state: GraphState) -> GraphState:
    vectorstore = state.get("vectorstore")
    sources = state.get("sources", [])
    if vectorstore is None:
        state["doc_context"] = ""
        state["retrieval_strong"] = False
        state["sources"] = sources
        return state
    try:
        docs, retrieval_strong = assess_retrieval_quality(
            vectorstore, state["retrieval_query"]
        )
        if docs:
            state["doc_context"] = "\n\n".join(d.page_content for d in docs)
            sources.extend(sorted({d.metadata.get("source", "document") for d in docs}))
        else:
            state["doc_context"] = ""
        state["retrieval_strong"] = retrieval_strong
        state["sources"] = sources
        logger.info(
            "Doc retrieval: %d chunks, strong=%s", len(docs), retrieval_strong
        )
    except Exception as e:
        logger.error("Document retrieval node failed: %s", e)
        state["doc_context"] = ""
        state["retrieval_strong"] = False
        state["sources"] = sources
    return state


def _retrieval_router(state: GraphState) -> str:
    if should_use_web_search(
        state.get("vectorstore"),
        state["retrieval_strong"],
        state["intent"],
        state["doc_context"],
    ):
        return "web_search_node"
    return "response_node"


def _web_search_node(state: GraphState) -> GraphState:
    sources = state.get("sources", [])
    try:
        web_context = gather_web_context(state["retrieval_query"])
        if web_context:
            state["web_context"] = web_context
            sources.append("web search")
            logger.info("Web search: %d chars", len(web_context))
        else:
            state["web_context"] = ""
    except Exception as e:
        logger.error("Web search node failed: %s", e)
        state["web_context"] = ""
    state["sources"] = sources
    return state


def _general_chat_node(state: GraphState) -> GraphState:
    """
    No-op pass-through node. Used for:
      - general_chat intent (greetings / identity)
      - follow_up_chat intent (memory-only, no docs/web needed)
    Memory context is already populated by _memory_node.
    """
    logger.info("General chat node: skipping retrieval and web search")
    return state


def _tabular_analytics_node(state: GraphState) -> GraphState:
    """
    Execute SQL or math on loaded DataFrames.

    On success: injects result as doc_context, sets retrieval_strong=True,
    routes to response_node.

    On failure: falls back to document_retrieval_node (schema injected as
    doc_context so the LLM has structural information).
    """
    sources = state.get("sources", [])
    try:
        result_text, success = run_tabular_analytics(
            state["question"], state["conversation_id"]
        )
        state["doc_context"] = result_text
        state["retrieval_strong"] = success
        if success:
            sources.append("tabular analytics")
            logger.info("Tabular analytics node: success, %d chars", len(result_text))
        else:
            logger.info(
                "Tabular analytics node: no mapping found, falling back to retrieval"
            )
    except Exception as e:
        logger.error("Tabular analytics node failed: %s", e)
        state["doc_context"] = ""
        state["retrieval_strong"] = False
    state["sources"] = sources
    return state


def _analytics_router(state: GraphState) -> str:
    if state["retrieval_strong"]:
        return "response_node"
    return "document_retrieval_node"


def _response_node(state: GraphState) -> GraphState:
    prompt = build_prompt(
        short_term_history=state["short_term"],
        long_term_summary=state["long_term_summary"],
        doc_context=state["doc_context"],
        web_context=state["web_context"],
        question=state["question"],
    )
    try:
        state["answer"] = state["llm"].invoke(prompt)
    except Exception as e:
        logger.error("LLM call failed in response_node: %s", e)
        raise
    logger.info("Response node: answer=%d chars", len(state["answer"]))
    return state


# ---------------------------------------------------------------------------
# Graph construction — compiled once at import time
# ---------------------------------------------------------------------------
def _build_graph() -> Any:
    workflow = StateGraph(GraphState)

    workflow.add_node("intent_classifier_node", _intent_classifier_node)
    workflow.add_node("memory_node", _memory_node)
    workflow.add_node("document_retrieval_node", _document_retrieval_node)
    workflow.add_node("web_search_node", _web_search_node)
    workflow.add_node("general_chat_node", _general_chat_node)
    workflow.add_node("tabular_analytics_node", _tabular_analytics_node)
    workflow.add_node("response_node", _response_node)

    workflow.set_entry_point("intent_classifier_node")
    workflow.add_edge("intent_classifier_node", "memory_node")

    workflow.add_conditional_edges(
        "memory_node",
        _memory_router,
        {
            "document_retrieval_node": "document_retrieval_node",
            "web_search_node": "web_search_node",
            "general_chat_node": "general_chat_node",
            "tabular_analytics_node": "tabular_analytics_node",
        },
    )

    workflow.add_conditional_edges(
        "tabular_analytics_node",
        _analytics_router,
        {
            "response_node": "response_node",
            "document_retrieval_node": "document_retrieval_node",
        },
    )

    workflow.add_conditional_edges(
        "document_retrieval_node",
        _retrieval_router,
        {
            "web_search_node": "web_search_node",
            "response_node": "response_node",
        },
    )

    workflow.add_edge("web_search_node", "response_node")
    workflow.add_edge("general_chat_node", "response_node")
    workflow.add_edge("response_node", END)

    return workflow.compile()


_graph = None


def _get_graph() -> Any:
    global _graph
    if _graph is None:
        _graph = _build_graph()
    return _graph


# ===========================================================================
# 17. Public Entry Point
# ===========================================================================
def generate_answer(
    llm: QwenNgrokLLM,
    conversation_id: str,
    vectorstore: Optional[FAISS],
    question: str,
) -> Tuple[str, List[str]]:
    """
    Main entry point called by backend_server.py.
    Runs the LangGraph pipeline; falls back to direct pipeline on graph failure.
    """
    initial_state: GraphState = {
        "question": question,
        "conversation_id": conversation_id,
        "vectorstore": vectorstore,
        "llm": llm,
        "intent": "",
        "retrieval_query": question,
        "is_follow_up": False,
        "short_term": "",
        "long_term_summary": "",
        "raw_messages": [],
        "doc_context": "",
        "retrieval_strong": False,
        "web_context": "",
        "answer": "",
        "sources": [],
    }

    try:
        final_state = _get_graph().invoke(initial_state)
        return final_state["answer"], final_state["sources"]
    except Exception as e:
        logger.error("Graph execution failed, running fallback pipeline: %s", e)
        return _fallback_pipeline(llm, conversation_id, vectorstore, question)


def _fallback_pipeline(
    llm: QwenNgrokLLM,
    conversation_id: str,
    vectorstore: Optional[FAISS],
    question: str,
) -> Tuple[str, List[str]]:
    """Direct linear pipeline — used only if LangGraph fails."""
    sources: List[str] = []

    try:
        short_term, long_term_summary, raw_messages = build_memory_context(
            conversation_id, question
        )
    except Exception:
        short_term, long_term_summary, raw_messages = "", "", []

    retrieval_query = (
        expand_query_with_context(question, raw_messages)
        if is_follow_up(question)
        else question
    )

    intent = _classify_intent(
        question, vectorstore, is_follow_up(question), conversation_id
    )

    doc_context = ""
    retrieval_strong = False

    if intent == "tabular_analytics":
        try:
            result_text, success = run_tabular_analytics(question, conversation_id)
            if success:
                doc_context = result_text
                retrieval_strong = True
                sources.append("tabular analytics")
        except Exception:
            pass

    if not doc_context:
        try:
            docs, retrieval_strong = assess_retrieval_quality(vectorstore, retrieval_query)
            if docs:
                doc_context = "\n\n".join(d.page_content for d in docs)
                sources.extend(
                    sorted({d.metadata.get("source", "document") for d in docs})
                )
        except Exception:
            pass

    web_context = ""
    if should_use_web_search(vectorstore, retrieval_strong, intent, doc_context):
        try:
            web_context = gather_web_context(retrieval_query)
            if web_context:
                sources.append("web search")
        except Exception:
            pass

    prompt = build_prompt(short_term, long_term_summary, doc_context, web_context, question)
    try:
        answer = llm.invoke(prompt)
    except Exception as e:
        logger.error("Fallback pipeline LLM call failed: %s", e)
        raise
    return answer, sources


# ===========================================================================
# Standalone self-test
# ===========================================================================
if __name__ == "__main__":
    logger.info("Backend self-test...")
    if check_llm_health():
        logger.info("LLM endpoint reachable.")
    else:
        logger.warning("LLM endpoint NOT reachable. Check NGROK_URL in .env.")